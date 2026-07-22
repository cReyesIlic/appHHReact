import asyncio
import json
import sqlite3
import tempfile
import unittest
from contextlib import closing
from io import BytesIO
from pathlib import Path
from types import SimpleNamespace
from unittest.mock import AsyncMock, Mock, patch

from docx import Document

from app.agents.agent_loop import AgentLoop, draft_tool_schemas
from app.agents.tools.handlers import search_rag
from app.core.config import settings
from app.services.proposal_drafts import ProposalDraftService


class ProposalDraftBriefTests(unittest.TestCase):
    def setUp(self):
        self.temp = tempfile.TemporaryDirectory()
        self.root = Path(self.temp.name)
        self.database_patch = patch.object(
            settings,
            "database_dir",
            str(self.root / "drafts.sqlite"),
        )
        self.database_patch.start()
        self.storage_patch = patch.object(
            ProposalDraftService,
            "_owner_root",
            lambda _service, owner_id: self.root / "proposal_drafts" / str(owner_id),
        )
        self.storage_patch.start()

    def tearDown(self):
        self.storage_patch.stop()
        self.database_patch.stop()
        self.temp.cleanup()

    def test_legacy_database_is_migrated_and_brief_invalidates_old_guide(self):
        with closing(sqlite3.connect(settings.sqlite_path)) as conn, conn:
            conn.execute(
                """
                create table proposal_drafts (
                    slug text primary key,
                    owner_id text not null,
                    title text not null,
                    cliente text,
                    status text default 'draft',
                    created_at text not null,
                    updated_at text not null
                )
                """
            )

        service = ProposalDraftService()
        draft = service.create_draft("owner@shimin.cl", "Nueva restitución", "Collahuasi")
        service.save_guide("owner@shimin.cl", draft["slug"], "# Guía antigua")

        updated = service.update_brief(
            "owner@shimin.cl",
            draft["slug"],
            "Preparar ingeniería de detalle y validar HH por disciplina.",
        )

        self.assertEqual(updated["status"], "draft")
        self.assertIn("validar HH", updated["brief_text"])
        self.assertFalse(updated["guide_exists"])
        with closing(sqlite3.connect(settings.sqlite_path)) as conn:
            columns = {row[1] for row in conn.execute("pragma table_info(proposal_drafts)")}
        self.assertIn("brief_text", columns)

    def test_uploaded_docx_is_extracted_and_searchable(self):
        service = ProposalDraftService()
        draft = service.create_draft("owner@shimin.cl", "Sistema de bombeo")
        document = Document()
        document.add_heading("Requisitos del cliente", level=1)
        document.add_paragraph("El plazo contractual es de doce semanas y exige un informe final.")
        buffer = BytesIO()
        document.save(buffer)

        result = service.add_file(
            "owner@shimin.cl",
            draft["slug"],
            "Bases técnicas.docx",
            buffer.getvalue(),
        )
        hits = service.search_chunks(draft["slug"], "plazo contractual informe", limit=5)

        self.assertGreater(result["chars_extracted"], 20)
        self.assertGreater(result["chunks_created"], 0)
        self.assertTrue(hits)
        self.assertIn("doce semanas", hits[0]["snippet"])

    def test_reupload_replaces_metadata_and_chunks_instead_of_duplicating(self):
        service = ProposalDraftService()
        draft = service.create_draft("owner@shimin.cl", "Truck Shop")

        def document_bytes(text: str) -> bytes:
            document = Document()
            document.add_paragraph(text)
            buffer = BytesIO()
            document.save(buffer)
            return buffer.getvalue()

        service.add_file(
            "owner@shimin.cl",
            draft["slug"],
            "Bases.docx",
            document_bytes("Versión antigua sobre oficinas."),
        )
        service.add_file(
            "owner@shimin.cl",
            draft["slug"],
            "Bases.docx",
            document_bytes("Versión vigente para revisión del Truck Shop."),
        )

        self.assertEqual(len(service.list_files(draft["slug"])), 1)
        self.assertFalse(service.search_chunks(draft["slug"], "antigua oficinas"))
        self.assertTrue(service.search_chunks(draft["slug"], "vigente truck"))

    def test_delete_file_removes_asset_text_chunks_and_invalidates_guide(self):
        service = ProposalDraftService()
        draft = service.create_draft("owner@shimin.cl", "Barrio Cívico")
        document = Document()
        document.add_paragraph("Informe de revisión de ingeniería de detalles.")
        buffer = BytesIO()
        document.save(buffer)
        service.add_file("owner@shimin.cl", draft["slug"], "Bases.docx", buffer.getvalue())
        service.save_guide("owner@shimin.cl", draft["slug"], "# Guía")

        result = service.delete_file("owner@shimin.cl", draft["slug"], "Bases.docx")

        self.assertTrue(result["deleted"])
        self.assertEqual(service.list_files(draft["slug"]), [])
        self.assertFalse(service.file_path("owner@shimin.cl", draft["slug"], "Bases.docx").exists())
        self.assertFalse(service.search_chunks(draft["slug"], "ingeniería detalles"))
        self.assertFalse(service.get_draft("owner@shimin.cl", draft["slug"])["guide_exists"])

    def test_document_register_and_review_estimate_use_review_rates(self):
        service = ProposalDraftService()
        draft = service.create_draft("owner@shimin.cl", "Revisión Truck Shop")
        document = Document()
        document.add_paragraph("3. Se comparten los siguientes documentos nuevos y actualizados.")
        document.add_paragraph("WOR-GPR-21CS187-OT-030-1128-E-DW-001 - 0")
        document.add_paragraph("WOR - GPR - 21CS187 - OT - 030 - 1128 - P - PID - 001 1 2")
        document.add_paragraph("CEN-ST-000-E-SK-003_v0 Estándar de puesta a tierra")
        document.add_paragraph(
            "WOR-GPR-21CS187-OT-030-1126G-C-DW-0126G-C-DW-006 texto OCR concatenado"
        )
        buffer = BytesIO()
        document.save(buffer)
        service.add_file("owner@shimin.cl", draft["slug"], "Registro.docx", buffer.getvalue())

        register = service.analyze_document_register("owner@shimin.cl", draft["slug"])
        estimate = service.estimate_document_review_hours(
            "owner@shimin.cl",
            draft["slug"],
            default_hours=5,
            hours_by_type={"PID": 8, "SK": 6},
            general_activities={"Informe final": 20},
            basis="Benchmark histórico de revisión",
        )

        self.assertEqual(register["total_documents"], 3)
        self.assertEqual(register["by_type"], {"DW": 1, "PID": 1, "SK": 1})
        self.assertEqual(estimate["review_hours"], 19)
        self.assertEqual(estimate["total_hours"], 39)
        self.assertEqual(len(estimate["document_rows"]), 3)

    def test_workspace_sections_are_persistent_and_overwrite_by_stage(self):
        service = ProposalDraftService()
        draft = service.create_draft("owner@shimin.cl", "Oferta progresiva")

        service.save_workspace_section(
            "owner@shimin.cl",
            draft["slug"],
            "hours",
            "Primera estimación",
            tables=[{"name": "HH", "rows": [{"total": 100}]}],
            sources=[{"title": "O-1553", "codigo": "O-1553"}],
        )
        service.save_workspace_section(
            "owner@shimin.cl",
            draft["slug"],
            "hours",
            "Estimación revisada",
            tables=[{"name": "HH", "rows": [{"total": 120}]}],
        )

        sections = service.get_draft("owner@shimin.cl", draft["slug"])["workspace_sections"]
        self.assertEqual(len(sections), 1)
        self.assertEqual(sections[0]["key"], "hours")
        self.assertEqual(sections[0]["content"], "Estimación revisada")
        self.assertEqual(sections[0]["tables"][0]["rows"][0]["total"], 120)

    def test_agent_checkpoint_persists_recipe_and_resumes_evidence(self):
        service = ProposalDraftService()
        draft = service.create_draft("owner@shimin.cl", "Oferta con checkpoint")
        started = service.start_agent_checkpoint(
            "owner@shimin.cl",
            draft["slug"],
            "hours",
            "Estimar revisión documental",
        )
        benchmark = {
            "codigo": "O-1553",
            "actividad": "Revisión 91 documentos",
            "documentos": 91,
            "hh": 455,
            "hh_por_documento": 5,
        }
        service.update_agent_checkpoint(
            "owner@shimin.cl",
            draft["slug"],
            {
                "status": "evidence_needed",
                "current_step": "quantify",
                "completed_steps": ["context", "inventory", "discover"],
                "completed_tools": ["search_rag", "get_hh_licitadas"],
                "quantitative_benchmarks": [benchmark],
                "evidence_gaps": ["Faltan 2 proyectos"],
            },
        )

        resumed = service.start_agent_checkpoint(
            "owner@shimin.cl",
            draft["slug"],
            "hours",
            "Continuar estimación",
        )
        checkpoint = service.get_draft("owner@shimin.cl", draft["slug"])["agent_checkpoint"]

        self.assertEqual(started["run_count"], 1)
        self.assertEqual(resumed["run_count"], 2)
        self.assertEqual(checkpoint["quantitative_benchmarks"], [benchmark])
        self.assertIn("get_hh_licitadas", checkpoint["completed_tools"])
        self.assertEqual(checkpoint["recipe"][0]["status"], "completed")

    def test_document_intelligence_ocr_preserves_page_evidence(self):
        service = ProposalDraftService()
        submit = Mock(status_code=202, headers={"operation-location": "https://ocr/jobs/1"})
        submit.raise_for_status = Mock()
        poll = Mock(
            headers={},
            json=Mock(
                return_value={
                    "status": "succeeded",
                    "analyzeResult": {
                        "pages": [
                            {
                                "pageNumber": 3,
                                "lines": [
                                    {"content": "Alcance del servicio"},
                                    {"content": "Plazo de doce semanas"},
                                ],
                            }
                        ]
                    },
                }
            ),
        )
        poll.raise_for_status = Mock()
        client = Mock()
        client.__enter__ = Mock(return_value=client)
        client.__exit__ = Mock(return_value=False)
        client.post.return_value = submit
        client.get.return_value = poll

        with patch.object(
            settings,
            "document_intelligence_endpoint",
            "https://docintel.example",
        ), patch.object(settings, "document_intelligence_key", "test-key"), patch(
            "app.services.proposal_drafts.httpx.Client",
            return_value=client,
        ):
            text = service._extract_pdf_with_document_intelligence(b"pdf-scanned")

        self.assertIn("[página 3]", text)
        self.assertIn("Plazo de doce semanas", text)
        self.assertEqual(client.post.call_args.kwargs["content"], b"pdf-scanned")


class ProposalDraftAgentContextTests(unittest.TestCase):
    def test_hours_checkpoint_requires_three_distinct_quantitative_projects(self):
        loop = AgentLoop.__new__(AgentLoop)
        updates = []
        state = {
            "completed_steps": ["context", "inventory", "discover", "estimate"],
            "quantitative_benchmarks": [
                {"codigo": "O-1553", "actividad": "Revisión A"},
                {"codigo": "O-1553", "actividad": "Revisión B"},
            ],
        }

        async def persist(**changes):
            updates.append(changes)

        asyncio.run(loop._finish_draft_checkpoint(persist, state, "hours", success=True))

        self.assertEqual(updates[-1]["status"], "evidence_needed")
        self.assertEqual(updates[-1]["current_step"], "quantify")
        self.assertIn("1/3", updates[-1]["evidence_gaps"][0])

        state["quantitative_benchmarks"].extend(
            [
                {"codigo": "O-1597", "actividad": "Revisión C"},
                {"codigo": "O-2410", "actividad": "Revisión D"},
            ]
        )
        asyncio.run(loop._finish_draft_checkpoint(persist, state, "hours", success=True))
        self.assertEqual(updates[-1]["status"], "completed")
        self.assertEqual(updates[-1]["evidence_gaps"], [])

    def test_quantitative_checkpoint_only_counts_review_rows_with_denominator(self):
        loop = AgentLoop.__new__(AgentLoop)
        result = {
            "codigo": "O-1553",
            "rows": [
                {
                    "key": "Revisión Electricidad (80 planos y 11 documentos)",
                    "total_hours": 455,
                },
                {"key": "Informe final", "total_hours": 39},
            ],
        }

        benchmarks = loop._extract_quantitative_review_benchmarks(
            "get_hh_licitadas",
            {"codigo": "O-1553"},
            result,
        )

        self.assertEqual(len(benchmarks), 1)
        self.assertEqual(benchmarks[0]["documentos"], 91)
        self.assertEqual(benchmarks[0]["hh_por_documento"], 5)
        self.assertEqual(loop._benchmark_project_count(benchmarks), 1)

    def test_table_deduplication_keeps_the_most_complete_version(self):
        loop = AgentLoop.__new__(AgentLoop)
        tables = [
            {"name": "Estimación por documento", "rows": []},
            {"name": "Otra tabla", "rows": [{"id": 1}]},
            {"name": "Estimación por documento", "rows": [{"id": 1}, {"id": 2}]},
        ]

        loop._dedupe_tables(tables)

        self.assertEqual([table["name"] for table in tables], ["Estimación por documento", "Otra tabla"])
        self.assertEqual(len(tables[0]["rows"]), 2)

    def test_rag_merges_multiple_semantic_queries(self):
        hybrid = Mock()

        async def search(query, filters, limit):
            common = {
                "codigo": "O-1553",
                "title": "Tabla de revisión",
                "summary": "HH para revisar documentos",
                "score": 2.0,
                "metadata": {},
            }
            if query == "constructibility review":
                return [common, {**common, "codigo": "O-1999", "title": "Constructibilidad"}]
            return [common]

        hybrid.search = AsyncMock(side_effect=search)
        result = asyncio.run(
            search_rag(
                SimpleNamespace(hybrid_rag=hybrid),
                queries=["revisión ingeniería", "constructibility review"],
                limit=8,
            )
        )

        self.assertEqual(result["queries_used"], ["revisión ingeniería", "constructibility review"])
        self.assertEqual(result["count"], 2)
        self.assertEqual(hybrid.search.await_count, 2)

    def test_draft_tools_are_agentic_but_expensive_costing_is_explicit(self):
        general = {
            schema["function"]["name"]
            for schema in draft_tool_schemas("Ayúdame a armar la revisión de ingeniería")
        }
        costing = {
            schema["function"]["name"]
            for schema in draft_tool_schemas("Estima las HH, costo y tarifa de O-1779")
        }

        self.assertTrue({"search_rag", "search_master", "search_draft_chunks"} <= general)
        self.assertNotIn("load_skill", general)
        self.assertNotIn("compute_economics", general)
        self.assertIn("compute_economics", costing)
        self.assertIn("get_hh_licitadas", costing)
        self.assertIn("get_proposal_detail", costing)

        hours_stage = {
            schema["function"]["name"]
            for schema in draft_tool_schemas("Estima la revisión", stage="hours")
        }
        self.assertNotIn("analyze_draft_document_register", hours_stage)
        self.assertIn("estimate_draft_review_hours", hours_stage)
        self.assertIn("read_pdf_deep", hours_stage)

    def test_active_draft_is_preloaded_before_the_model_answers(self):
        loop = AgentLoop.__new__(AgentLoop)
        loop.max_iterations = 1
        loop.skill_registry = Mock()
        loop.skill_registry.catalog.return_value = ""
        loop.llm = Mock(client=object(), azure=False)
        loop.llm.chat_with_tools = AsyncMock(
            return_value=SimpleNamespace(
                tool_calls=[],
                content="Sugiero estructurar alcance, entregables, riesgos y HH.",
            )
        )

        dispatcher = Mock()

        async def dispatch(name, args):
            if name == "load_skill":
                return {"name": "armar_propuesta", "content": "Flujo por etapas"}
            if name == "get_draft_context":
                return {
                    "slug": "restitucion-abc123",
                    "title": "Restitución",
                    "brief_text": "Necesito preparar la oferta.",
                    "chunks_preview": "El cliente solicita ingeniería de detalle.",
                    "source_assets": [
                        {
                            "title": "Bases técnicas.pdf",
                            "url": "/api/drafts/restitucion-abc123/files/Bases%20t%C3%A9cnicas.pdf",
                        }
                    ],
                }
            if name == "search_draft_chunks":
                return {
                    "count": 1,
                    "hits": [
                        {
                            "title": "Bases técnicas.pdf",
                            "url": "/api/drafts/restitucion-abc123/files/Bases%20t%C3%A9cnicas.pdf",
                            "snippet": "El cliente solicita ingeniería de detalle.",
                            "score": 2.0,
                        }
                    ],
                }
            raise AssertionError(name)

        dispatcher.dispatch = AsyncMock(side_effect=dispatch)
        with patch("app.agents.agent_loop.ToolContext.build", return_value=object()), patch(
            "app.agents.agent_loop.ToolDispatcher", return_value=dispatcher
        ):
            result = asyncio.run(
                loop.run(
                    "¿Cómo debería armar esta propuesta?",
                    active_draft={
                        "slug": "restitucion-abc123",
                        "title": "Restitución",
                        "cliente": "Collahuasi",
                        "brief_text": "Necesito preparar la oferta.",
                    },
                )
            )

        self.assertEqual(
            [call.args[0] for call in dispatcher.dispatch.await_args_list],
            ["load_skill", "get_draft_context", "search_draft_chunks"],
        )
        model_messages = loop.llm.chat_with_tools.await_args.kwargs["messages"]
        payload = json.loads(model_messages[-1]["content"])
        self.assertEqual(payload["draft_activo"]["slug"], "restitucion-abc123")
        self.assertIn("contexto_draft", payload)
        self.assertIn("/api/drafts/restitucion-abc123/files/", result.answer)
        self.assertTrue(any(source.url for source in result.sources))

    def test_active_draft_times_out_slow_tool_and_still_closes_answer(self):
        loop = AgentLoop.__new__(AgentLoop)
        loop.max_iterations = 1
        loop.skill_registry = Mock()
        loop.skill_registry.catalog.return_value = ""
        tool_call = SimpleNamespace(
            id="call-1",
            function=SimpleNamespace(
                name="compute_proposal_support",
                arguments=json.dumps({"query": "revisión ingeniería", "limit": 5}),
            ),
        )
        loop.llm = Mock(client=object(), azure=False)
        loop.llm.chat_with_tools = AsyncMock(
            side_effect=[
                SimpleNamespace(tool_calls=[tool_call], content=""),
                SimpleNamespace(tool_calls=[], content="Propuesta accionable cerrada."),
            ]
        )
        dispatcher = Mock()

        async def dispatch(name, args):
            if name == "load_skill":
                return {"name": "armar_propuesta", "content": "Flujo"}
            if name in {"get_draft_context", "search_draft_chunks"}:
                return {"slug": "revision-id", "count": 0}
            if name == "compute_proposal_support":
                await asyncio.sleep(0.1)
                return {"referencias_directas": []}
            raise AssertionError(name)

        dispatcher.dispatch = AsyncMock(side_effect=dispatch)
        with patch("app.agents.agent_loop.ToolContext.build", return_value=object()), patch(
            "app.agents.agent_loop.ToolDispatcher", return_value=dispatcher
        ), patch("app.agents.agent_loop.DRAFT_TOOL_TIMEOUT_SECONDS", 0.01):
            result = asyncio.run(
                loop.run(
                    "Ayúdame a armar esta propuesta",
                    active_draft={"slug": "revision-id", "title": "Revisión", "brief_text": "Alcance"},
                )
            )

        self.assertEqual(result.answer, "Propuesta accionable cerrada.")
        self.assertTrue(
            any(t.tool == "agent.compute_proposal_support" and t.status == "error" for t in result.trace)
        )

    def test_hours_stage_forces_ai_estimator_before_final_answer(self):
        loop = AgentLoop.__new__(AgentLoop)
        loop.max_iterations = 1
        loop.skill_registry = Mock()
        loop.skill_registry.catalog.return_value = ""
        evidence_call = SimpleNamespace(
            id="call-evidence",
            function=SimpleNamespace(
                name="search_rag",
                arguments=json.dumps({"query": "revisión ingeniería", "limit": 4}),
            ),
        )
        estimate_call = SimpleNamespace(
            id="call-estimate",
            function=SimpleNamespace(
                name="estimate_draft_review_hours",
                arguments=json.dumps(
                    {
                        "slug": "revision-id",
                        "default_hours": 5,
                        "hours_by_type": {"PID": 8},
                        "general_activities": {"Informe final": 20},
                        "basis": "O-1553: revisión histórica por documento",
                    }
                ),
            ),
        )
        loop.llm = Mock(client=object(), azure=False)
        loop.llm.chat_with_tools = AsyncMock(
            side_effect=[
                SimpleNamespace(tool_calls=[evidence_call], content=""),
                SimpleNamespace(tool_calls=[estimate_call], content=""),
                SimpleNamespace(tool_calls=[], content="## Estimación\nTotal: **39 HH**"),
            ]
        )
        dispatcher = Mock()

        async def dispatch(name, args):
            if name == "load_skill":
                return {"name": "armar_propuesta", "content": "Flujo"}
            if name in {"get_draft_context", "search_draft_chunks"}:
                return {"slug": "revision-id", "count": 0}
            if name == "analyze_draft_document_register":
                return {"total_documents": 3, "tables": []}
            if name == "search_rag":
                return {"count": 1, "hits": [{"codigo": "O-1553", "summary": "5 HH/documento"}]}
            if name == "estimate_draft_review_hours":
                return {
                    "total_documents": 3,
                    "review_hours": 19,
                    "general_hours": 20,
                    "total_hours": 39,
                    "tables": [
                        {"name": "Estimación por documento a revisar", "rows": [{"documento": "A", "hh": 5}]}
                    ],
                }
            raise AssertionError(name)

        dispatcher.dispatch = AsyncMock(side_effect=dispatch)
        with patch("app.agents.agent_loop.ToolContext.build", return_value=object()), patch(
            "app.agents.agent_loop.ToolDispatcher", return_value=dispatcher
        ):
            result = asyncio.run(
                loop.run(
                    "Estima las HH de revisión",
                    active_draft={
                        "slug": "revision-id",
                        "title": "Revisión",
                        "brief_text": "Revisar documentos, no elaborarlos",
                        "stage": "hours",
                    },
                )
            )

        called_names = [call.args[0] for call in dispatcher.dispatch.await_args_list]
        self.assertEqual(called_names.count("analyze_draft_document_register"), 1)
        self.assertIn("estimate_draft_review_hours", called_names)
        self.assertEqual(
            loop.llm.chat_with_tools.await_args_list[1].kwargs["tool_choice"],
            {"type": "function", "function": {"name": "estimate_draft_review_hours"}},
        )
        self.assertIn("39 HH", result.answer)
        self.assertTrue(any(table.get("name") == "Estimación por documento a revisar" for table in result.tables))
        self.assertTrue(
            any(
                item.tool == "agent.estimate_draft_review_hours" and item.status == "ok"
                for item in result.trace
            )
        )


if __name__ == "__main__":
    unittest.main()
