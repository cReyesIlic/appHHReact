import tempfile
import unittest
from pathlib import Path
from unittest.mock import patch

from docx import Document
from PyPDF2 import PdfReader

from app.core.config import settings
from app.schemas import ExportRequest, Source
from app.services.exports import ExportService, _parse_markdown


MARKDOWN_SAMPLE = """# Detalle de propuesta O-1779

## Ficha comercial

**Estado:** PG — propuesta ganada.
*Cliente:* Collahuasi. Usa `895 HH` como total.

[Abrir PDF en SharePoint](https://sharepoint.example/O-1779.pdf)

- HH licitadas por entregable
- Alcance documental verificado

1. Consultar Excel.
2. Validar PDF.

| Entregable | HH | Roles |
|---|---:|---|
| Informe final | **52** | JP 8, CN 2, ESP 42 |
| Visita a terreno | 10 | ESP 10 |

> Fuente: Master y documento emitido.

```text
O-1779 / control de trazabilidad
```
"""


class ExportMarkdownTests(unittest.TestCase):
    def request(self) -> ExportRequest:
        return ExportRequest(
            title="Detalle O-1779",
            answer=MARKDOWN_SAMPLE,
            sources=[
                Source(
                    kind="sharepoint",
                    title="O-1779.pdf",
                    codigo="O-1779",
                    url="https://sharepoint.example/O-1779.pdf",
                )
            ],
        )

    def test_parser_recognizes_markdown_table_and_rich_blocks(self):
        blocks = _parse_markdown(MARKDOWN_SAMPLE)
        kinds = [block.kind for block in blocks]
        self.assertIn("heading", kinds)
        self.assertIn("bullet", kinds)
        self.assertIn("number", kinds)
        self.assertIn("table", kinds)
        self.assertIn("quote", kinds)
        self.assertIn("code", kinds)
        table = next(block for block in blocks if block.kind == "table")
        self.assertEqual(table.headers, ["Entregable", "HH", "Roles"])
        self.assertEqual(len(table.rows), 2)

    def test_docx_renders_formatting_tables_and_links_without_raw_markdown(self):
        with tempfile.TemporaryDirectory() as temp_dir, patch.object(
            settings, "export_dir", Path(temp_dir)
        ):
            path = ExportService().create("docx", self.request())
            document = Document(path)
            text = "\n".join(
                [paragraph.text for paragraph in document.paragraphs]
                + [cell.text for table in document.tables for row in table.rows for cell in row.cells]
            )
            xml = document.part._element.xml
            relationships = " ".join(str(rel.target_ref) for rel in document.part.rels.values())

        self.assertNotIn("**", text)
        self.assertNotIn("|---", text)
        self.assertNotIn("](https://", text)
        self.assertIn("Detalle de propuesta O-1779", text)
        self.assertIn("Informe final", text)
        self.assertGreaterEqual(len(document.tables), 1)
        self.assertIn("w:hyperlink", xml)
        self.assertIn("https://sharepoint.example/O-1779.pdf", relationships)

    def test_pdf_and_branded_report_render_markdown_as_document_content(self):
        with tempfile.TemporaryDirectory() as temp_dir, patch.object(
            settings, "export_dir", Path(temp_dir)
        ):
            pdf_path = ExportService().create("pdf", self.request())
            pdf_text = "\n".join(page.extract_text() or "" for page in PdfReader(pdf_path).pages)
            report_path = ExportService().create("report", self.request())
            report_text = "\n".join(page.extract_text() or "" for page in PdfReader(report_path).pages)

        for rendered in (pdf_text, report_text):
            self.assertNotIn("**", rendered)
            self.assertNotIn("|---", rendered)
            self.assertNotIn("](https://", rendered)
            self.assertIn("Detalle de propuesta O-1779", rendered)
            self.assertIn("Informe final", rendered)
            self.assertIn("52", rendered)


if __name__ == "__main__":
    unittest.main()
