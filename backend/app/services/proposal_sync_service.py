"""Sincronización end-to-end SharePoint → Master/RAG/Wiki.

Flujo por código:
  1. Lista PDFs en SharePoint (último válido)
  2. Descarga PDF localmente (cache)
  3. Parsea texto completo + páginas
  4. Indexa en `rag_parent_sections + rag_child_chunks` (tabla agente híbrido)
  5. Genera embeddings (`HybridRagStore.build()` con `--force=False` solo indexa pendientes)
  6. Compila página Wiki vía `WikiAutoCompiler.compile_for_proposal()`
  7. Actualiza manifest CSV

Idempotente:
  - Si el código ya tiene parent_sections, replace_chunks lo reemplaza.
  - Si la página Wiki ya existe y `force_wiki=False`, skip.
  - Si los embeddings ya existen para el child_id+model, skip (LEFT JOIN filter).

Modo dry-run:
  - `discover_new()` lista códigos en SharePoint que NO están aún en master/RAG/wiki.
  - No descarga, no llama LLM, no indexa.
"""

from __future__ import annotations

import asyncio
import csv
import json
import os
import sqlite3
import time
from contextlib import closing
from dataclasses import dataclass, field
from datetime import datetime
from pathlib import Path

from app.core.config import settings
from app.rag.hybrid_store import HybridRagStore
from app.rag.parent_child import ParentChildIndexer
from app.services.email_client import EmailClient
from app.services.ingestion_reporter import ganadas_sync_report
from app.services.knowledge_extractor import KnowledgeExtractor
from app.services.knowledge_models import ProposalMetadata
from app.services.master_repository import MasterRepository
from app.services.pipeline_registry import (
    PIPELINE_VERSION,
    RAG_PIPELINE_VERSION,
    WIKI_PIPELINE_VERSION,
    PipelineRegistry,
    source_signature,
)
from app.services.proposal_taxonomy import enrich_metadata
from app.services.rag_store import RagStore
from app.services.sharepoint_client import SharePointClient, normalize_offer_code
from app.services.structured_wiki import StructuredWikiService
from app.services.wiki_auto_compiler import WikiAutoCompiler


MANIFEST_PATH = "storage/sync_manifest.csv"
MANIFEST_COLUMNS = [
    "codigo",
    "status",
    "pdf_name",
    "chunks_parent",
    "chunks_child",
    "wiki_status",
    "wiki_path",
    "error",
    "updated_at",
]


@dataclass
class SyncCounters:
    discovered: int = 0
    ingested: int = 0
    skipped: int = 0
    errors: int = 0
    wiki_ok: int = 0
    wiki_skipped: int = 0
    wiki_no_rag: int = 0
    wiki_error: int = 0
    by_code: list[dict] = field(default_factory=list)


class ProposalSyncService:
    def __init__(self) -> None:
        self.sharepoint = SharePointClient()
        self.master = MasterRepository()
        self.extractor = KnowledgeExtractor()
        self.rag_store = RagStore()
        self.parent_child = ParentChildIndexer()
        self.hybrid = HybridRagStore()
        self.wiki = StructuredWikiService()
        self.wiki_compiler = WikiAutoCompiler()
        self.pipeline = PipelineRegistry()

    # ---- discovery ----

    def discover_ganadas_pendientes(self, include_excel: bool = True) -> dict:
        """Detecta propuestas GANADAS (PG) en master que aún NO están indexadas en RAG.

        Este es el flujo correcto post-adjudicación:
        - Master se actualiza con el estado PG cuando SHIMIN gana una propuesta.
        - Para esas ganadas hay que traer los PDFs (y opcionalmente Excel) desde SharePoint
          e indexarlos para que el agente las use como referencias confiables.

        NO mira todo SharePoint, solo lo que el master marcó como ganado.
        """
        ganadas = self._ganadas_master_rows()

        already = self._codigos_with_rag()
        states = self.pipeline.by_codes([row["codigo"] for row in ganadas])
        proposals_dir = settings.resolve_path("storage/llm_wiki/proposals")
        wiki_existing = {p.stem.upper() for p in proposals_dir.glob("O-*.md")} if proposals_dir.exists() else set()

        pendientes_rag = self._order_pending_by_last_attempt(
            [{**g, "sync_reason": "new_rag"} for g in ganadas if g["codigo"] not in already]
        )
        pendientes_reprocess = self._order_pending_by_last_attempt(
            [
                {**g, "sync_reason": "pipeline_stale"}
                for g in ganadas
                if g["codigo"] in already
                and (
                    not states.get(g["codigo"])
                    or states[g["codigo"]].get("pipeline_version") != PIPELINE_VERSION
                    or states[g["codigo"]].get("rag_pipeline_version") != RAG_PIPELINE_VERSION
                    or states[g["codigo"]].get("wiki_pipeline_version") != WIKI_PIPELINE_VERSION
                    or states[g["codigo"]].get("needs_reprocess")
                )
            ]
        )
        pendientes_wiki = [g for g in ganadas if g["codigo"] not in wiki_existing]
        total_ganadas = len(ganadas)

        return {
            "total_ganadas_master": total_ganadas,
            "ya_indexadas_rag": total_ganadas - len(pendientes_rag),
            "ya_compiladas_wiki": total_ganadas - len(pendientes_wiki),
            "pendientes_rag_count": len(pendientes_rag),
            "pendientes_reprocess_count": len(pendientes_reprocess),
            "pendientes_wiki_count": len(pendientes_wiki),
            "pipeline_version": PIPELINE_VERSION,
            # La cola completa es necesaria para que el scheduler pueda rotarla.
            # El limite de cada corrida se aplica en sync_ganadas(), no aqui.
            "pendientes_rag": pendientes_rag,
            "pendientes_reprocess": pendientes_reprocess,
            "pendientes_wiki": pendientes_wiki,
        }

    async def sync_ganadas(self, limit: int = 20, include_excel: bool = True) -> dict:
        """Para cada propuesta ganada (PG) en master sin RAG indexado, descarga PDFs y Excel
        de SharePoint y la alimenta al sistema (RAG + Wiki + embeddings).

        Idempotente. Apto para uso bajo demanda (botón UI o cron manual cuando se confirme).
        """
        gap = self.discover_ganadas_pendientes(include_excel=include_excel)
        limit = max(1, int(limit))
        stale = gap.get("pendientes_reprocess") or []
        pending_codes = {row["codigo"] for row in gap["pendientes_rag"]}
        stale_codes = {row["codigo"] for row in stale}
        recheck_candidates = [
            row for row in self._ganadas_master_rows()
            if row["codigo"] not in pending_codes and row["codigo"] not in stale_codes
        ]
        changed = await self._discover_changed_sources(recheck_candidates)
        pendientes = self._mix_queues([changed, gap["pendientes_rag"], stale], limit)
        counters = SyncCounters(discovered=len(pendientes))
        for ganada in pendientes:
            codigo = ganada["codigo"]
            reason = ganada.get("sync_reason") or "new_rag"
            outcome = await self.sync_code(
                codigo,
                force_wiki=reason != "new_rag",
                source_files=ganada.get("_source_files"),
            )
            outcome["sync_reason"] = reason
            counters.by_code.append({**outcome, "titulo": ganada.get("titulo"), "cliente": ganada.get("cliente")})
            if outcome["status"] == "ok":
                counters.ingested += 1
            elif outcome["status"] in {"skipped", "no_pdf", "no_files"}:
                counters.skipped += 1
            else:
                counters.errors += 1
            if outcome.get("wiki_status") == "ok":
                counters.wiki_ok += 1
            elif outcome.get("wiki_status") == "no_rag":
                counters.wiki_no_rag += 1
            elif outcome.get("wiki_status") == "error":
                counters.wiki_error += 1
        summary = {
            "scope": "ganadas_master_pendientes",
            "total_ganadas_master": gap["total_ganadas_master"],
            "ya_indexadas": gap["ya_indexadas_rag"],
            "pipeline_version": PIPELINE_VERSION,
            "sources_changed": len(changed),
            "pipeline_stale": len(stale),
            "objetivo_corrida": len(pendientes),
            "ingested": counters.ingested,
            "skipped": counters.skipped,
            "errors": counters.errors,
            "wiki_ok": counters.wiki_ok,
            "wiki_no_rag": counters.wiki_no_rag,
            "wiki_error": counters.wiki_error,
            "details": counters.by_code,
        }
        # Reporte por email — best effort, no rompe el sync si falla
        try:
            email = EmailClient()
            if email.configured:
                subject, text, html = ganadas_sync_report(summary, kind="ganadas")
                send_result = email.send(subject, text, html)
                summary["email"] = send_result
            else:
                summary["email"] = {"sent": False, "reason": "ACS no configurado"}
        except Exception as exc:  # noqa: BLE001
            summary["email"] = {"sent": False, "reason": f"{type(exc).__name__}: {exc}"}
        return summary

    async def discover_new(self, limit: int = 200) -> dict:
        """Lista códigos en SharePoint que aún NO están indexados en RAG parent_child.

        (Legacy — barre TODO SharePoint, no discrimina por estado. Para uso normal,
        prefiere `discover_ganadas_pendientes()` que es más certero.)
        """
        folders = await self.sharepoint.list_offer_folders(limit=limit)
        already = self._codigos_with_rag()
        new = [f for f in folders if f.get("codigo", "").upper() not in already]
        return {
            "sharepoint_total": len(folders),
            "already_indexed": len(folders) - len(new),
            "new_count": len(new),
            "new": new[:limit],
        }

    def discover_wiki_gaps(self, only_with_rag: bool = True) -> dict:
        """Devuelve códigos con RAG indexado pero SIN página Wiki en disco."""
        rag_codes = self._codigos_with_rag()
        proposals_dir = settings.resolve_path("storage/llm_wiki/proposals")
        existing = {p.stem.upper() for p in proposals_dir.glob("O-*.md")} if proposals_dir.exists() else set()
        missing = sorted(rag_codes - existing)
        return {
            "rag_count": len(rag_codes),
            "wiki_pages": len(existing),
            "missing_wiki": len(missing),
            "missing_codes": missing,
        }

    # ---- sync per code ----

    async def sync_code(
        self,
        codigo: str,
        *,
        force_wiki: bool = False,
        build_embeddings: bool = True,
        source_files: list[dict] | None = None,
    ) -> dict:
        codigo = codigo.upper().strip()
        result = {
            "codigo": codigo,
            "status": "pending",
            "chunks_parent": 0,
            "chunks_child": 0,
            "wiki_status": "skipped",
        }
        previous_state = self.pipeline.get(codigo)
        previous_source_codes = self._source_offer_codes((previous_state or {}).get("source_files") or [])
        previous_source_mismatch = bool(previous_source_codes and codigo not in previous_source_codes)
        if previous_source_mismatch:
            # La Wiki anterior tampoco se puede reutilizar si fue construida con
            # documentos de otra propuesta.
            force_wiki = True

        # 1-3. Descargar y parsear (PDF, DOCX, XLSX)
        try:
            files = source_files if source_files is not None else await self.sharepoint.list_emitido_files(codigo)
            if not files:
                # Fallback: intentar la lógica antigua que solo busca PDFs
                pdfs = await self.sharepoint.list_pdfs(codigo)
                latest = self.sharepoint.select_latest_pdf(pdfs)
                if not latest:
                    if previous_source_mismatch:
                        mismatch = ", ".join(sorted(previous_source_codes))
                        error = f"Indice retirado: las fuentes registradas pertenecian a {mismatch}, no a {codigo}"
                        cleanup = self._purge_invalid_index(codigo, previous_state or {})
                        result.update({"status": "invalid_source_removed", "error": error, "cleanup": cleanup})
                        self._record_manifest(result)
                        self.pipeline.record_invalidated(codigo, error)
                        return result
                    result.update({"status": "no_files", "note": "carpeta '03 Oferta/02 Emitido' vacía o sin PDF/DOCX/XLSX"})
                    self._record_manifest(result)
                    self.pipeline.record_failure(codigo, result)
                    return result
                files = [latest]
            current_source_codes = self._source_offer_codes(files)
            if current_source_codes and codigo not in current_source_codes:
                mismatch = ", ".join(sorted(current_source_codes))
                error = f"Fuente rechazada: corresponde a {mismatch}, no a {codigo}"
                cleanup = None
                if previous_source_mismatch:
                    cleanup = self._purge_invalid_index(codigo, previous_state or {})
                    self.pipeline.record_invalidated(codigo, error)
                else:
                    self.pipeline.record_failure(codigo, {"status": "source_mismatch", "error": error})
                result.update({"status": "source_mismatch", "error": error})
                if cleanup is not None:
                    result["cleanup"] = cleanup
                self._record_manifest(result)
                return result
            # Procesar todos los archivos emitidos → concatenar texto
            full_text_parts = []
            processed_files = []
            file_errors = []
            excel_assets = []
            first_pages_text = ""
            primary_name = None
            primary_path = None
            primary_url = None
            for f in files:
                content = await self.sharepoint.download_file(f)
                local_path = self.sharepoint.save_pdf_locally(
                    codigo,
                    f.get("name", "doc.bin"),
                    content,
                    item_id=f.get("id"),
                )
                kind = (f.get("name", "").lower().rsplit(".", 1)[-1] if "." in f.get("name", "") else "")
                file_text = self._extract_text_any(content, kind, f.get("name", ""))
                if not file_text.strip() or file_text.startswith("[error extrayendo "):
                    file_errors.append(file_text or f"{f.get('name')}: sin texto extraible")
                    continue
                processed_files.append(f)
                if kind in {"xlsx", "xlsm", "xls"}:
                    excel_assets.append(
                        {
                            "name": f.get("name") or local_path.name,
                            "source_file": local_path.name,
                            "kind": kind,
                            "content": content,
                            "local_path": str(local_path),
                        }
                    )
                full_text_parts.append(f"\n\n## Fuente: {f.get('name')}\n\n{file_text}")
                # primer archivo procesable = "primario" (para metadata + name)
                if primary_name is None:
                    primary_name = f.get("name")
                    primary_path = str(local_path)
                    primary_url = f.get("webUrl")
                    if kind == "pdf":
                        first_pages_text = self.sharepoint.extract_first_pages_text(content, pages=5)
                    else:
                        first_pages_text = file_text[:8000]
            full_text = "\n".join(full_text_parts)
            if not full_text.strip():
                result.update(
                    {
                        "status": "no_text",
                        "error": "Los archivos emitidos no contienen texto extraible",
                        "file_errors": file_errors,
                    }
                )
                self._record_manifest(result)
                self.pipeline.record_failure(codigo, result)
                return result
            result["text_chars"] = len(full_text)
            latest = {"name": primary_name, "webUrl": primary_url}
            local_path = primary_path or ""
            result["pdf_name"] = primary_name
            result["files_discovered"] = len(files)
            result["files_processed"] = len(processed_files)
            result["file_errors"] = file_errors
            result["kinds_processed"] = sorted(
                set(f.get("name", "").lower().rsplit(".", 1)[-1] for f in processed_files if "." in f.get("name", ""))
            )
            if excel_assets:
                excel_processing = await self._process_excel_assets(codigo, excel_assets)
                result["excel_processing"] = excel_processing
                result["excel_parsed_files"] = sum(
                    1 for row in excel_processing if row.get("hh_status") == "ok" or row.get("budget_persisted")
                )
                result["excel_errors"] = [
                    row["error"] for row in excel_processing if row.get("error")
                ]
            result["source_signature"] = source_signature(files)
        except Exception as exc:  # noqa: BLE001
            result.update({"status": "error", "error": f"sharepoint: {exc}"})
            self._record_manifest(result)
            self.pipeline.record_failure(codigo, result)
            return result

        # 4. Indexar en parent_child (tabla híbrida)
        metadata = self._metadata(codigo, latest, str(local_path))
        try:
            knowledge = await self.extractor.extract(metadata, first_pages_text, full_text)
            raw_metadata = {**metadata.model_dump(), **knowledge.model_dump()}
            enriched = enrich_metadata(raw_metadata)
            parse_result = {"text": full_text, "pages": []}
            pc_result = self.parent_child.index_parse_result(codigo, parse_result, enriched)
            result["chunks_parent"] = pc_result.get("parents", 0)
            result["chunks_child"] = pc_result.get("children", 0)
            # Legacy chunks (rag_chunks) — opcional, mantiene compat
            chunks = self.rag_store.make_chunks(
                codigo=codigo,
                text=full_text,
                source=latest.get("webUrl") or str(local_path),
                metadata=raw_metadata,
            )
            self.rag_store.upsert_proposal(metadata, knowledge)
            self.rag_store.replace_chunks(codigo, chunks)
        except Exception as exc:  # noqa: BLE001
            result.update({"status": "error", "error": f"indexing: {exc}"})
            self._record_manifest(result)
            self.pipeline.record_failure(codigo, result)
            return result

        # 5. Embeddings (solo pendientes — idempotente)
        if build_embeddings:
            try:
                # Solo procesa chunks de este codigo si están sin embedding
                embedding_result = await self._build_embeddings_for_code(codigo)
                result["embedding_count"] = int(embedding_result.get("processed") or 0)
            except Exception as exc:  # noqa: BLE001
                result["embedding_error"] = str(exc)

        # 6. Wiki autocompile
        try:
            wiki_res = await self.wiki_compiler.compile_for_proposal(codigo, force=force_wiki)
            result["wiki_status"] = wiki_res.get("status")
            result["wiki_path"] = wiki_res.get("path")
            result["wiki_entry_id"] = wiki_res.get("entry_id")
            result["quality"] = self._quality_result(result, wiki_res.get("quality") or {})
        except Exception as exc:  # noqa: BLE001
            result["wiki_status"] = "error"
            result["wiki_error"] = str(exc)
            result["quality"] = self._quality_result(result, {})

        result["status"] = "ok"
        self._record_manifest(result)
        self.pipeline.record_success(codigo, files, result)
        return result

    async def sync_new(self, limit: int = 50, force_wiki: bool = False) -> dict:
        discovery = await self.discover_new(limit=limit)
        counters = SyncCounters(discovered=discovery["new_count"])
        for folder in discovery["new"]:
            codigo = folder.get("codigo")
            if not codigo:
                continue
            outcome = await self.sync_code(codigo, force_wiki=force_wiki)
            counters.by_code.append(outcome)
            if outcome["status"] == "ok":
                counters.ingested += 1
            elif outcome["status"] in {"skipped", "no_pdf", "no_files"}:
                counters.skipped += 1
            else:
                counters.errors += 1
            wiki_st = outcome.get("wiki_status")
            if wiki_st == "ok":
                counters.wiki_ok += 1
            elif wiki_st == "skipped":
                counters.wiki_skipped += 1
            elif wiki_st == "no_rag":
                counters.wiki_no_rag += 1
            elif wiki_st in {"error"}:
                counters.wiki_error += 1
        return {
            "discovered": counters.discovered,
            "ingested": counters.ingested,
            "skipped": counters.skipped,
            "errors": counters.errors,
            "wiki_ok": counters.wiki_ok,
            "wiki_skipped": counters.wiki_skipped,
            "wiki_no_rag": counters.wiki_no_rag,
            "wiki_error": counters.wiki_error,
            "details": counters.by_code,
        }

    # ---- backfill wiki for already-RAG'd proposals ----

    async def backfill_wiki(
        self,
        codigos: list[str] | None = None,
        *,
        force: bool = False,
        limit: int | None = None,
        concurrency: int = 8,
    ) -> dict:
        """Compila páginas Wiki para códigos que ya tienen RAG indexado, en paralelo.

        Si `codigos` es None, usa `discover_wiki_gaps`. Útil para 1500+ propuestas existentes.
        Concurrencia por defecto = 8 (ajustable; Azure tiene rate limits — bajar si fallan llamadas).
        """
        import asyncio

        if codigos:
            target = [c.upper() for c in codigos]
        else:
            gap = self.discover_wiki_gaps()
            target = gap["missing_codes"]
        if limit:
            target = target[:limit]

        counters = SyncCounters(discovered=len(target))
        sem = asyncio.Semaphore(max(1, concurrency))
        progress = {"done": 0}
        total = len(target)
        progress_lock = asyncio.Lock()

        async def worker(codigo: str) -> dict:
            async with sem:
                try:
                    outcome = await self.wiki_compiler.compile_for_proposal(codigo, force=force)
                except Exception as exc:  # noqa: BLE001
                    outcome = {"codigo": codigo, "status": "error", "error": str(exc)}
                async with progress_lock:
                    progress["done"] += 1
                    if progress["done"] % 25 == 0 or progress["done"] == total:
                        print(
                            f"  · backfill: {progress['done']}/{total} ok={counters.wiki_ok} err={counters.wiki_error}",
                            flush=True,
                        )
                if outcome.get("status") == "ok":
                    counters.wiki_ok += 1
                elif outcome.get("status") == "skipped":
                    counters.wiki_skipped += 1
                elif outcome.get("status") == "no_rag":
                    counters.wiki_no_rag += 1
                else:
                    counters.wiki_error += 1
                counters.by_code.append(outcome)
                return outcome

        await asyncio.gather(*(worker(c) for c in target))
        return {
            "target_count": total,
            "wiki_ok": counters.wiki_ok,
            "wiki_skipped": counters.wiki_skipped,
            "wiki_no_rag": counters.wiki_no_rag,
            "wiki_error": counters.wiki_error,
            "details": counters.by_code,
        }

    # ---- internal helpers ----

    async def _process_excel_assets(self, codigo: str, assets: list[dict]) -> list[dict]:
        """Extrae HH localmente y, si esta configurado, normaliza presupuesto en paralelo."""
        from app.services.budget_extractor_client import BudgetExtractorClient
        from app.services.hh_excel_extractor import HHExcelExtractor

        hh = HHExcelExtractor()
        budget = BudgetExtractorClient()
        results: list[dict] = []
        for asset in assets:
            row = {
                "name": asset["name"],
                "source_file": asset["source_file"],
                "hh_status": "unsupported",
                "hh_rows": 0,
            }
            if asset.get("kind") in {"xlsx", "xlsm"}:
                hh_result = hh.extract_file(codigo, asset["local_path"])
                row["hh_status"] = hh_result.get("status")
                row["hh_rows"] = int(hh_result.get("rows") or 0)
                if hh_result.get("status") == "error":
                    row["error"] = f"{asset['name']}: HH: {hh_result.get('error') or 'error'}"
            results.append(row)

        if not budget.available:
            return results

        extracted = await asyncio.gather(
            *(
                budget.extract_normalized(codigo, asset["content"], asset["source_file"])
                for asset in assets
            )
        )
        for row, payload in zip(results, extracted):
            if payload.get("error"):
                error = f"{row['name']}: presupuesto: {payload['error']}"
                row["error"] = f"{row['error']}; {error}" if row.get("error") else error
                continue
            row["budget_persisted"] = budget.persist(codigo, row["source_file"], payload)
            row["budget_totals"] = payload.get("totals")
        return results

    def _source_offer_codes(self, files: list[dict]) -> set[str]:
        explicit = {
            str(row.get("sourceOfferCode") or row.get("source_offer_code") or "").upper()
            for row in files
            if row.get("sourceOfferCode") or row.get("source_offer_code")
        }
        if explicit:
            return explicit
        detected: set[str] = set()
        for row in files:
            # La URL de los PDFs incluye la carpeta de oferta. No usamos el
            # nombre del archivo: una oferta valida puede citar otra propuesta.
            value = str(row.get("webUrl") or row.get("web_url") or "")
            code = normalize_offer_code(value)
            if code:
                detected.add(code)
        return detected

    def _purge_invalid_index(self, codigo: str, state: dict) -> dict:
        """Retira RAG/Wiki contaminados, limitado al codigo comprobado."""
        codigo = codigo.upper()
        removed: dict[str, int | bool | str] = {}
        tables = (
            "rag_child_embeddings",
            "rag_child_chunks",
            "rag_parent_sections",
            "rag_chunks",
            "proposal_knowledge",
        )
        with closing(sqlite3.connect(settings.sqlite_path, timeout=30)) as conn, conn:
                existing = {
                    row[0]
                    for row in conn.execute("select name from sqlite_master where type='table'").fetchall()
                }
                for table in tables:
                    if table not in existing:
                        removed[table] = 0
                        continue
                    count = conn.execute(f"select count(*) from {table} where codigo = ?", (codigo,)).fetchone()[0]
                    conn.execute(f"delete from {table} where codigo = ?", (codigo,))
                    removed[table] = int(count or 0)

        entry_id = str(state.get("wiki_entry_id") or "").strip()
        if entry_id:
            try:
                self.wiki.delete_entry(entry_id)
                removed["wiki_entry"] = entry_id
            except KeyError:
                removed["wiki_entry"] = False
        proposal_page = settings.resolve_path(f"storage/llm_wiki/proposals/{codigo}.md")
        if proposal_page.exists():
            proposal_page.unlink()
            removed["wiki_page"] = True
        else:
            removed["wiki_page"] = False
        cache_dir = settings.resolve_path(f"storage/proposals/{codigo}")
        cache_files_removed = 0
        if cache_dir.exists():
            for path in cache_dir.iterdir():
                if path.is_file():
                    path.unlink()
                    cache_files_removed += 1
            if not any(cache_dir.iterdir()):
                cache_dir.rmdir()
        removed["cache_files"] = cache_files_removed
        StructuredWikiService.invalidate_sync_cache()
        return removed

    def _ganadas_master_rows(self) -> list[dict]:
        from app.services.proposal_taxonomy import status_category

        unique: dict[str, dict] = {}
        for row in self.master.all_offers():
            estado = str(row.get("estado") or "").strip().upper()
            codigo = str(row.get("codigo") or "").strip().upper()
            if status_category(estado) != "ganada" or not codigo.startswith("O-"):
                continue
            unique[codigo] = {
                "codigo": codigo,
                "titulo": str(row.get("titulo") or "")[:160],
                "cliente": row.get("cliente_directo") or row.get("cliente_final"),
                "fecha_recep": row.get("fecha_recep") or row.get("fecha_recepcion"),
                "estado": estado,
                "monto": row.get("monto"),
                "horas_lic": row.get("horas_lic"),
                "cod_proy": row.get("cod_proy"),
            }
        return list(unique.values())

    async def _discover_changed_sources(self, candidates: list[dict]) -> list[dict]:
        """Revisa en rotacion fuentes ya procesadas y detecta cambios por firma Graph."""
        if not candidates:
            return []
        states = self.pipeline.by_codes([row["codigo"] for row in candidates])
        ordered = sorted(
            [row for row in candidates if states.get(row["codigo"])],
            key=lambda row: str(states[row["codigo"]].get("source_checked_at") or ""),
        )
        recheck_limit = self._bounded_env_int("SYNC_SOURCE_RECHECK_LIMIT", 200, minimum=1, maximum=5000)
        concurrency = self._bounded_env_int("SYNC_SOURCE_RECHECK_CONCURRENCY", 8, minimum=1, maximum=20)
        semaphore = asyncio.Semaphore(concurrency)

        async def inspect(row: dict) -> dict | None:
            codigo = row["codigo"]
            async with semaphore:
                try:
                    files = await self.sharepoint.list_emitido_files(codigo)
                except Exception:
                    return None
            if not files:
                return None
            stored = states[codigo].get("source_signature") or ""
            current = source_signature(files)
            self.pipeline.mark_checked(codigo, files, establish_baseline=False)
            if stored and current != stored:
                return {
                    **row,
                    "sync_reason": "source_changed",
                    "_source_files": files,
                }
            return None

        results = await asyncio.gather(*(inspect(row) for row in ordered[:recheck_limit]))
        return [row for row in results if row]

    def _mix_queues(self, queues: list[list[dict]], limit: int) -> list[dict]:
        """Round-robin entre cambios, propuestas nuevas y versiones obsoletas."""
        pending = [list(queue) for queue in queues]
        selected: list[dict] = []
        seen: set[str] = set()
        while len(selected) < limit and any(pending):
            for queue in pending:
                while queue:
                    row = queue.pop(0)
                    codigo = str(row.get("codigo") or "").upper()
                    if not codigo or codigo in seen:
                        continue
                    selected.append(row)
                    seen.add(codigo)
                    break
                if len(selected) >= limit:
                    break
        return selected

    def _quality_result(self, result: dict, ai_quality: dict) -> dict:
        rag_score = 0
        rag_score += 15 if int(result.get("files_processed") or 0) > 0 else 0
        rag_score += 20 if int(result.get("text_chars") or 0) >= 5000 else 8 if result.get("text_chars") else 0
        rag_score += 20 if int(result.get("chunks_parent") or 0) > 0 else 0
        rag_score += 20 if int(result.get("chunks_child") or 0) > 0 else 0
        rag_score += 25 if int(result.get("embedding_count") or 0) > 0 and not result.get("embedding_error") else 0
        if result.get("file_errors"):
            rag_score = max(0, rag_score - 15)
        if result.get("excel_errors"):
            rag_score = max(0, rag_score - 10)
        wiki_score = ai_quality.get("wiki_score")
        if wiki_score is None:
            wiki_score = 70 if result.get("wiki_status") == "ok" else 45 if result.get("wiki_status") == "skipped" else 0
        return {
            "mode": ai_quality.get("mode") or "heuristic",
            "rag_score": ai_quality.get("rag_score", rag_score),
            "wiki_score": wiki_score,
            "summary": ai_quality.get("summary") or "Calidad calculada con señales objetivas del pipeline.",
            "issues": [
                *(ai_quality.get("issues") or []),
                *result.get("file_errors", []),
                *result.get("excel_errors", []),
            ][:8],
        }

    def _bounded_env_int(self, name: str, default: int, *, minimum: int, maximum: int) -> int:
        try:
            value = int(os.getenv(name, str(default)))
        except (TypeError, ValueError):
            value = default
        return max(minimum, min(maximum, value))

    def _codigos_with_rag(self) -> set[str]:
        with closing(sqlite3.connect(settings.sqlite_path, timeout=5)) as conn:
            rows = conn.execute("select distinct codigo from rag_parent_sections").fetchall()
        return {str(r[0]).upper() for r in rows if r[0]}

    def _order_pending_by_last_attempt(self, pending: list[dict]) -> list[dict]:
        """Ordena la cola por el intento mas antiguo, dejando nuevos primero.

        Cada resultado se agrega al manifest. Al mover lo recien intentado al
        final, un codigo sin archivos o con error no puede bloquear para siempre
        las propuestas que vienen despues.
        """
        attempts = self._latest_manifest_attempts()

        def key(row: dict) -> str:
            codigo = str(row.get("codigo") or "").upper()
            return attempts.get(codigo, {}).get("updated_at") or ""

        return sorted(pending, key=key)

    def _latest_manifest_attempts(self) -> dict[str, dict]:
        path = settings.resolve_path(MANIFEST_PATH)
        if not path.exists():
            return {}
        latest: dict[str, dict] = {}
        try:
            with path.open("r", encoding="utf-8-sig", newline="") as fh:
                for row in csv.DictReader(fh):
                    codigo = str(row.get("codigo") or "").strip().upper()
                    updated_at = str(row.get("updated_at") or "").strip()
                    if not codigo:
                        continue
                    previous = latest.get(codigo)
                    if previous is None or updated_at >= str(previous.get("updated_at") or ""):
                        latest[codigo] = row
        except (OSError, csv.Error):
            return {}
        return latest

    async def _build_embeddings_for_code(self, codigo: str) -> dict:
        """Genera embeddings solo para los chunks del código sin embedding (del modelo activo)."""
        # HybridRagStore.build() ya filtra por LEFT JOIN sobre embeddings del modelo activo,
        # pero no por codigo. Hacemos una corrida limitada apuntada con un fetch directo.
        model = self.hybrid.embeddings.deployment
        with closing(sqlite3.connect(settings.sqlite_path, timeout=5)) as conn:
            conn.row_factory = sqlite3.Row
            rows = [
                dict(r)
                for r in conn.execute(
                    """
                    select c.child_id, c.parent_id, c.codigo, c.text, c.page_start, c.page_end, c.metadata,
                           p.title, p.text as parent_text, p.metadata as parent_metadata
                    from rag_child_chunks c
                    join rag_parent_sections p on p.parent_id = c.parent_id
                    left join rag_child_embeddings e on e.child_id = c.child_id and e.model = ?
                    where c.codigo = ? and e.child_id is null
                    order by c.parent_id, c.child_id
                    """,
                    (model, codigo.upper()),
                ).fetchall()
            ]
        if not rows:
            return {"selected": 0, "processed": 0}
        for row in rows:
            metadata = json.loads(row.get("metadata") or "{}")
            parent_metadata = json.loads(row.get("parent_metadata") or "{}")
            row["metadata_dict"] = metadata
            row["parent_metadata_dict"] = parent_metadata
            row["embedding_text"] = self.hybrid._embedding_text(row, metadata, parent_metadata)
            row["content_hash"] = self.hybrid.embeddings.content_hash(row["embedding_text"])
        vectors = await self.hybrid.embeddings.embed_texts([r["embedding_text"] for r in rows])
        self.hybrid._save_batch(rows, vectors)
        return {"selected": len(rows), "processed": len(rows)}

    async def ingest_local_file(self, codigo: str, content: bytes, filename: str) -> dict:
        """Ingesta un archivo subido manualmente: PDF/DOCX → RAG; XLSX → RAG + HH extractor.

        Guarda en storage/emitted_offer_assets/{pdf|excel}/{codigo}/{filename},
        actualiza manifest.csv de assets, e indexa en parent_child + embeddings.
        """
        import re
        from app.services.hh_excel_extractor import HHExcelExtractor

        codigo = (codigo or "").upper().strip().replace(" ", "")
        if not re.match(r"^[OS]H?-?\d{2,6}$", codigo) or any(c in codigo for c in ("/", "\\", "..", "\x00")):
            return {"codigo": codigo, "status": "invalid_codigo", "error": f"codigo invalido: {codigo!r}"}
        kind = filename.lower().rsplit(".", 1)[-1] if "." in filename else ""
        if kind not in {"pdf", "docx", "xlsx", "xls", "xlsm"}:
            return {"codigo": codigo, "status": "unsupported", "error": f"Formato no soportado: {kind}"}

        safe_name = re.sub(r"[^A-Za-z0-9._ -]+", "_", filename).strip() or "archivo.bin"
        subdir = "excel" if kind in {"xlsx", "xls", "xlsm"} else "pdf"
        target = settings.resolve_path(f"storage/emitted_offer_assets/{subdir}/{codigo}/{safe_name}")
        target.parent.mkdir(parents=True, exist_ok=True)
        target.write_bytes(content)

        result = {
            "codigo": codigo,
            "status": "pending",
            "filename": safe_name,
            "kind": kind,
            "saved_to": str(target),
        }

        text = self._extract_text_any(content, kind, filename)
        if not text or text.startswith("[error"):
            result.update({"status": "extract_error", "error": text or "sin texto"})
            self._update_assets_manifest(codigo, kind, safe_name, target)
            return result

        try:
            virtual = {"name": safe_name, "webUrl": None}
            metadata = self._metadata(codigo, virtual, str(target))
            knowledge = await self.extractor.extract(metadata, text[:8000], text)
            raw_metadata = {**metadata.model_dump(), **knowledge.model_dump()}
            enriched = enrich_metadata(raw_metadata)
            parse_result = {"text": text, "pages": []}
            pc = self.parent_child.index_parse_result(codigo, parse_result, enriched)
            result["chunks_parent"] = pc.get("parents", 0)
            result["chunks_child"] = pc.get("children", 0)
            chunks = self.rag_store.make_chunks(codigo=codigo, text=text, source=str(target), metadata=raw_metadata)
            self.rag_store.upsert_proposal(metadata, knowledge)
            self.rag_store.replace_chunks(codigo, chunks)
            try:
                await self._build_embeddings_for_code(codigo)
            except Exception as exc:  # noqa: BLE001
                result["embedding_error"] = str(exc)
        except Exception as exc:  # noqa: BLE001
            result.update({"status": "index_error", "error": f"indexing: {exc}"})
            self._update_assets_manifest(codigo, kind, safe_name, target)
            return result

        if kind in {"xlsx", "xlsm"}:
            try:
                hh_res = HHExcelExtractor().extract_file(codigo, target)
                result["hh_rows"] = hh_res.get("rows", 0)
                result["hh_status"] = hh_res.get("status")
            except Exception as exc:  # noqa: BLE001
                result["hh_error"] = str(exc)

            # Azure Function: extracción robusta (proyectos_extracted + tarifas + gastos)
            try:
                from app.services.budget_extractor_client import BudgetExtractorClient
                bclient = BudgetExtractorClient()
                if bclient.available:
                    extracted = await bclient.extract_normalized(codigo, content, safe_name)
                    if extracted.get("error"):
                        result["budget_extractor_error"] = extracted["error"]
                    else:
                        persisted = bclient.persist(codigo, safe_name, extracted)
                        result["budget_extractor"] = {
                            "totals": extracted.get("totals"),
                            "persisted": persisted,
                            "sheet": extracted.get("entregables_sheet_seleccionado"),
                        }
            except Exception as exc:  # noqa: BLE001
                result["budget_extractor_error"] = f"{type(exc).__name__}: {exc}"

        self._update_assets_manifest(codigo, kind, safe_name, target)
        result["status"] = "ok"
        return result

    def _update_assets_manifest(self, codigo: str, kind: str, filename: str, local_path: Path) -> None:
        """Actualiza (o crea) la fila del manifest emitted_offer_assets/manifest.csv para este codigo."""
        manifest_path = settings.resolve_path("storage/emitted_offer_assets/manifest.csv")
        manifest_path.parent.mkdir(parents=True, exist_ok=True)
        cols = ["codigo", "status", "folder_name", "pdf_count", "excel_count", "zip_count",
                "selected_pdf", "selected_pdf_local", "selected_excel", "selected_excel_local",
                "zip_assets", "error", "updated_at"]
        rows: list[dict] = []
        existing: dict | None = None
        if manifest_path.exists():
            with manifest_path.open("r", encoding="utf-8-sig", newline="") as fh:
                for r in csv.DictReader(fh):
                    if (r.get("codigo") or "").strip().upper() == codigo:
                        existing = r
                    else:
                        rows.append(r)
        row = existing or {c: "" for c in cols}
        row["codigo"] = codigo
        row["updated_at"] = datetime.now().isoformat(timespec="seconds")
        row["status"] = "manual_upload"
        if kind in {"pdf"}:
            try:
                row["pdf_count"] = str(int(row.get("pdf_count") or 0) + 1)
            except ValueError:
                row["pdf_count"] = "1"
            row["selected_pdf"] = filename
            row["selected_pdf_local"] = str(local_path)
        elif kind in {"xlsx", "xls", "xlsm"}:
            try:
                row["excel_count"] = str(int(row.get("excel_count") or 0) + 1)
            except ValueError:
                row["excel_count"] = "1"
            row["selected_excel"] = filename
            row["selected_excel_local"] = str(local_path)
        rows.append(row)
        with manifest_path.open("w", encoding="utf-8", newline="") as fh:
            writer = csv.DictWriter(fh, fieldnames=cols)
            writer.writeheader()
            for r in rows:
                writer.writerow({c: r.get(c, "") for c in cols})

    def _extract_text_any(self, content: bytes, kind: str, filename: str) -> str:
        """Extrae texto de PDF, DOCX o XLSX. Devuelve texto plano concatenado."""
        from io import BytesIO
        try:
            if kind == "pdf":
                from PyPDF2 import PdfReader
                reader = PdfReader(BytesIO(content))
                return "\n".join(p.extract_text() or "" for p in reader.pages)
            if kind == "docx":
                from docx import Document
                doc = Document(BytesIO(content))
                blocks = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
                for table in doc.tables:
                    for row in table.rows:
                        cells = [c.text.strip() for c in row.cells if c.text.strip()]
                        if cells:
                            blocks.append(" | ".join(cells))
                return "\n".join(blocks)
            if kind in {"xlsx", "xlsm"}:
                import openpyxl
                wb = openpyxl.load_workbook(BytesIO(content), data_only=True, read_only=True)
                blocks = []
                for ws in wb.worksheets:
                    blocks.append(f"## Hoja: {ws.title}")
                    for row in ws.iter_rows(values_only=True):
                        cells = [str(c).strip() for c in row if c is not None and str(c).strip()]
                        if cells:
                            blocks.append(" | ".join(cells))
                return "\n".join(blocks)
            if kind == "xls":
                import pandas as pd
                sheets = pd.read_excel(BytesIO(content), sheet_name=None, header=None)
                blocks = []
                for title, frame in sheets.items():
                    blocks.append(f"## Hoja: {title}")
                    for row in frame.fillna("").astype(str).itertuples(index=False, name=None):
                        cells = [cell.strip() for cell in row if cell.strip()]
                        if cells:
                            blocks.append(" | ".join(cells))
                return "\n".join(blocks)
        except Exception as exc:  # noqa: BLE001
            return f"[error extrayendo {filename}: {exc}]"
        return ""

    def _metadata(self, codigo: str, pdf: dict, local_path: str) -> ProposalMetadata:
        master = self.master.search(codigo=codigo, limit=1)
        row = master[0] if master else {}
        return ProposalMetadata(
            codigo=codigo,
            pdf_name=pdf.get("name", ""),
            url=pdf.get("webUrl"),
            local_path=local_path,
            cliente=row.get("cliente_directo"),
            cliente_final=row.get("cliente_final"),
            titulo=row.get("titulo"),
            estado=row.get("estado"),
            tipo_servicio=row.get("tipo_servicio"),
            fecha_recepcion=row.get("fecha_recep") or row.get("fecha_recepcion"),
        )

    def _record_manifest(self, result: dict) -> None:
        path = settings.resolve_path(MANIFEST_PATH)
        path.parent.mkdir(parents=True, exist_ok=True)
        is_new = not path.exists()
        with path.open("a", encoding="utf-8", newline="") as fh:
            writer = csv.DictWriter(fh, fieldnames=MANIFEST_COLUMNS)
            if is_new:
                writer.writeheader()
            writer.writerow(
                {
                    "codigo": result.get("codigo", ""),
                    "status": result.get("status", ""),
                    "pdf_name": result.get("pdf_name", ""),
                    "chunks_parent": result.get("chunks_parent", 0),
                    "chunks_child": result.get("chunks_child", 0),
                    "wiki_status": result.get("wiki_status", ""),
                    "wiki_path": result.get("wiki_path", ""),
                    "error": result.get("error", ""),
                    "updated_at": datetime.now().isoformat(timespec="seconds"),
                }
            )
