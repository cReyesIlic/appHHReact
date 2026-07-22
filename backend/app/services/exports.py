from pathlib import Path
from datetime import datetime
from dataclasses import dataclass, field
from html import escape as html_escape
import re
import subprocess

import pandas as pd
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from reportlab.lib.pagesizes import letter, A4
from reportlab.lib import colors
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import cm
from reportlab.pdfgen import canvas
from reportlab.platypus import (
    SimpleDocTemplate,
    Paragraph,
    Spacer,
    Table,
    TableStyle,
    Image as RLImage,
    PageBreak,
)
from reportlab.platypus.flowables import HRFlowable

from app.core.config import settings
from app.schemas import ExportRequest


# Paleta SHIMIN
SHIMIN_INK = colors.HexColor("#182b36")
SHIMIN_DEEP = colors.HexColor("#223f4d")
SHIMIN_COPPER = colors.HexColor("#c8863b")
SHIMIN_GOLD = colors.HexColor("#e1b35f")
SHIMIN_MIST = colors.HexColor("#eef3f4")
SHIMIN_LINE = colors.HexColor("#d7e0e2")
SHIMIN_MUTED = colors.HexColor("#5f747d")
LOGO_PATH = Path(__file__).resolve().parents[1] / "assets" / "logo-shimin.png"


class ExportService:
    def create(self, kind: str, request: ExportRequest) -> Path:
        self.export_dir.mkdir(parents=True, exist_ok=True)
        if kind == "xlsx":
            return self._xlsx(request)
        if kind == "docx":
            return self._docx(request)
        if kind in {"typst", "typst-pdf"}:
            return self._typst_pdf(request)
        if kind in {"pdf", "report"}:
            return self._pdf(request)
        raise ValueError("Formato soportado: xlsx, docx, pdf")

    def _xlsx(self, request: ExportRequest) -> Path:
        path = self.export_dir / "respuesta.xlsx"
        generated_at = datetime.now().strftime("%d-%m-%Y %H:%M")
        with pd.ExcelWriter(path, engine="xlsxwriter") as writer:
            workbook = writer.book

            # Estilos SHIMIN
            title_fmt = workbook.add_format({
                "bold": True, "font_size": 18, "font_color": "#182b36",
                "font_name": "Calibri", "valign": "vcenter",
            })
            subtitle_fmt = workbook.add_format({
                "italic": True, "font_size": 10, "font_color": "#5f747d",
                "font_name": "Calibri",
            })
            section_fmt = workbook.add_format({
                "bold": True, "font_size": 12, "font_color": "white",
                "bg_color": "#182b36", "font_name": "Calibri", "align": "left",
                "valign": "vcenter", "border": 0,
            })
            body_fmt = workbook.add_format({
                "font_size": 10, "font_name": "Calibri", "text_wrap": True,
                "valign": "top",
            })
            header_fmt = workbook.add_format({
                "bold": True, "font_size": 10, "font_color": "white",
                "bg_color": "#c8863b", "font_name": "Calibri",
                "align": "center", "valign": "vcenter", "border": 1,
                "border_color": "#a96a23",
            })
            cell_fmt = workbook.add_format({
                "font_size": 9, "font_name": "Calibri",
                "border": 1, "border_color": "#d7e0e2", "valign": "top",
                "text_wrap": True,
            })
            alt_cell_fmt = workbook.add_format({
                "font_size": 9, "font_name": "Calibri",
                "border": 1, "border_color": "#d7e0e2", "valign": "top",
                "text_wrap": True, "bg_color": "#eef3f4",
            })

            # Portada
            cover = workbook.add_worksheet("Portada")
            cover.hide_gridlines(2)
            cover.set_column("A:A", 4)
            cover.set_column("B:B", 80)
            if LOGO_PATH.exists():
                cover.insert_image("B2", str(LOGO_PATH), {"x_scale": 0.6, "y_scale": 0.6})
            cover.set_row(8, 32)
            cover.write("B9", request.title or "Respuesta SHIMIN", title_fmt)
            cover.write("B10", f"Generado: {generated_at}  ·  SHIMIN Proposal Intelligence", subtitle_fmt)
            cover.set_row(12, 22)
            cover.write("B13", "Resumen", section_fmt)
            cover.set_column("B:B", 95)
            cover.write("B15", request.answer[:30000], body_fmt)
            cover.set_row(14, 380)

            # Tablas (cada una en su hoja)
            for idx, table in enumerate(request.tables[:10], start=1):
                rows = table.get("rows", [])
                if not rows:
                    continue
                name = str(table.get("name", f"Tabla {idx}"))[:31]
                ws = writer.book.add_worksheet(name)
                ws.hide_gridlines(2)
                columns = list(rows[0].keys())[:12]
                ws.set_row(0, 22)
                for col_idx, col in enumerate(columns):
                    ws.write(0, col_idx, str(col), header_fmt)
                    ws.set_column(col_idx, col_idx, max(12, min(40, len(str(col)) + 8)))
                for row_idx, row in enumerate(rows[:300], start=1):
                    fmt = alt_cell_fmt if row_idx % 2 == 0 else cell_fmt
                    for col_idx, col in enumerate(columns):
                        val = row.get(col, "")
                        ws.write(row_idx, col_idx, str(val)[:600], fmt)
                ws.freeze_panes(1, 0)
                ws.autofilter(0, 0, min(len(rows), 300), len(columns) - 1)

            # Fuentes
            if request.sources:
                ws = writer.book.add_worksheet("Fuentes")
                ws.hide_gridlines(2)
                headers = ["Tipo", "Código", "Título", "URL", "Score"]
                for col_idx, h in enumerate(headers):
                    ws.write(0, col_idx, h, header_fmt)
                ws.set_column(0, 0, 14)
                ws.set_column(1, 1, 12)
                ws.set_column(2, 2, 60)
                ws.set_column(3, 3, 40)
                ws.set_column(4, 4, 10)
                for row_idx, src in enumerate(request.sources[:200], start=1):
                    s = src.model_dump() if hasattr(src, "model_dump") else dict(src)
                    fmt = alt_cell_fmt if row_idx % 2 == 0 else cell_fmt
                    ws.write(row_idx, 0, str(s.get("kind", "")), fmt)
                    ws.write(row_idx, 1, str(s.get("codigo", "") or ""), fmt)
                    ws.write(row_idx, 2, str(s.get("title", ""))[:300], fmt)
                    ws.write(row_idx, 3, str(s.get("url", "") or ""), fmt)
                    ws.write(row_idx, 4, str(s.get("score", "") or ""), fmt)
                ws.freeze_panes(1, 0)
        return path

    def _docx(self, request: ExportRequest) -> Path:
        path = self.export_dir / "respuesta.docx"
        doc = Document()

        # Márgenes
        for section in doc.sections:
            section.left_margin = Cm(1.8)
            section.right_margin = Cm(1.8)
            section.top_margin = Cm(1.8)
            section.bottom_margin = Cm(1.8)

        # Estilos base
        styles = doc.styles
        normal = styles["Normal"]
        normal.font.name = "Calibri"
        normal.font.size = Pt(10.5)

        # Header con logo
        if LOGO_PATH.exists():
            try:
                header = doc.sections[0].header
                h_para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
                h_run = h_para.add_run()
                h_run.add_picture(str(LOGO_PATH), width=Inches(0.7))
                h_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            except Exception:
                pass

        # Portada — banda azul SHIMIN
        title_para = doc.add_paragraph()
        title_para.paragraph_format.space_before = Pt(20)
        title_para.paragraph_format.space_after = Pt(4)
        run = title_para.add_run(request.title or "Respuesta SHIMIN")
        run.bold = True
        run.font.size = Pt(22)
        run.font.color.rgb = RGBColor(0x18, 0x2B, 0x36)

        sub = doc.add_paragraph()
        sub_run = sub.add_run(
            f"SHIMIN Proposal Intelligence  ·  Generado {datetime.now().strftime('%d-%m-%Y %H:%M')}"
        )
        sub_run.italic = True
        sub_run.font.size = Pt(9.5)
        sub_run.font.color.rgb = RGBColor(0x5F, 0x74, 0x7D)

        # Línea separadora cobre
        _docx_horizontal_rule(doc, color="C8863B", size=8)

        # Cuerpo Markdown: títulos, énfasis, enlaces, listas, tablas y código.
        _docx_render_markdown(doc, request.answer or "")

        # Tablas
        for table in request.tables[:10]:
            rows = table.get("rows", [])
            if not rows:
                continue
            doc.add_paragraph()  # espacio
            heading = doc.add_paragraph()
            h_run = heading.add_run(str(table.get("name", "Tabla")))
            h_run.bold = True
            h_run.font.size = Pt(13)
            h_run.font.color.rgb = RGBColor(0x18, 0x2B, 0x36)
            heading.paragraph_format.space_before = Pt(14)
            heading.paragraph_format.space_after = Pt(4)

            columns = list(rows[0].keys())[:8]
            doc_table = doc.add_table(rows=1, cols=len(columns))
            doc_table.style = "Light Grid Accent 1"
            doc_table.autofit = True
            hdr_cells = doc_table.rows[0].cells
            for index, column in enumerate(columns):
                hdr_cells[index].text = ""
                cell_para = hdr_cells[index].paragraphs[0]
                cell_run = cell_para.add_run(str(column))
                cell_run.bold = True
                cell_run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                cell_run.font.size = Pt(9.5)
                _docx_shade_cell(hdr_cells[index], "C8863B")
            for r_idx, row in enumerate(rows[:60]):
                cells = doc_table.add_row().cells
                for c_idx, column in enumerate(columns):
                    cells[c_idx].text = _display_value(row.get(column, ""))[:500]
                    for paragraph in cells[c_idx].paragraphs:
                        for run in paragraph.runs:
                            run.font.size = Pt(9)
                    if r_idx % 2 == 0:
                        _docx_shade_cell(cells[c_idx], "EEF3F4")

        # Fuentes
        if request.sources:
            doc.add_paragraph()
            h = doc.add_paragraph()
            run = h.add_run("Fuentes")
            run.bold = True
            run.font.size = Pt(13)
            run.font.color.rgb = RGBColor(0x18, 0x2B, 0x36)
            h.paragraph_format.space_before = Pt(14)
            for source in request.sources[:30]:
                s = source.model_dump() if hasattr(source, "model_dump") else dict(source)
                line = doc.add_paragraph(style="List Bullet")
                kind_run = line.add_run(f"[{s.get('kind', 'src')}] ")
                kind_run.bold = True
                kind_run.font.size = Pt(9)
                kind_run.font.color.rgb = RGBColor(0xC8, 0x86, 0x3B)
                rest_text = s.get("title") or s.get("codigo") or ""
                if s.get("codigo"):
                    rest_text = f"{s['codigo']}  ·  {rest_text}"
                line.add_run(str(rest_text)[:280]).font.size = Pt(9.5)
                if s.get("url"):
                    line.add_run("  ")
                    _docx_add_hyperlink(line, "Abrir fuente", str(s["url"]))

        # Footer
        try:
            footer = doc.sections[0].footer
            f_para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
            f_run = f_para.add_run("SHIMIN Proposal Intelligence  ·  Confidencial")
            f_run.font.size = Pt(8)
            f_run.font.color.rgb = RGBColor(0x5F, 0x74, 0x7D)
            f_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        except Exception:
            pass

        doc.save(path)
        return path

    def _pdf(self, request: ExportRequest) -> Path:
        path = self.export_dir / "respuesta.pdf"
        doc = SimpleDocTemplate(
            str(path),
            pagesize=A4,
            leftMargin=1.8 * cm,
            rightMargin=1.8 * cm,
            topMargin=2.4 * cm,
            bottomMargin=2.0 * cm,
            title=request.title[:90],
            author="SHIMIN",
        )
        styles = getSampleStyleSheet()

        title_style = ParagraphStyle(
            "TitleS", parent=styles["Title"],
            fontName="Helvetica-Bold", fontSize=20, leading=24,
            textColor=SHIMIN_INK, spaceAfter=6,
        )
        sub_style = ParagraphStyle(
            "SubS", parent=styles["Normal"],
            fontName="Helvetica-Oblique", fontSize=9,
            textColor=SHIMIN_MUTED, spaceAfter=8,
        )
        body_style = ParagraphStyle(
            "BodyS", parent=styles["BodyText"],
            fontName="Helvetica", fontSize=10, leading=14,
            textColor=colors.HexColor("#202124"), spaceAfter=6,
        )
        h2_style = ParagraphStyle(
            "H2S", parent=styles["Heading2"],
            fontName="Helvetica-Bold", fontSize=14, leading=18,
            textColor=SHIMIN_DEEP, spaceBefore=12, spaceAfter=6,
        )
        h3_style = ParagraphStyle(
            "H3S", parent=styles["Heading3"],
            fontName="Helvetica-Bold", fontSize=11, leading=15,
            textColor=SHIMIN_COPPER, spaceBefore=8, spaceAfter=4,
        )
        bullet_style = ParagraphStyle(
            "BulletS", parent=body_style, leftIndent=14, firstLineIndent=-8,
            bulletIndent=5, spaceAfter=3,
        )
        quote_style = ParagraphStyle(
            "QuoteS", parent=body_style, leftIndent=12, rightIndent=8,
            borderColor=SHIMIN_COPPER, borderWidth=0.8, borderPadding=6,
            backColor=SHIMIN_MIST, textColor=SHIMIN_MUTED,
        )
        code_style = ParagraphStyle(
            "CodeS", parent=body_style, fontName="Courier", fontSize=8.5,
            leading=11, leftIndent=8, rightIndent=8, borderPadding=6,
            backColor=colors.HexColor("#f4f6f7"),
        )
        table_header_style = ParagraphStyle(
            "TableHeaderS", parent=body_style, fontName="Helvetica-Bold",
            fontSize=8.5, leading=10, textColor=colors.white, spaceAfter=0,
        )

        story = []
        # Portada con logo
        if LOGO_PATH.exists():
            try:
                story.append(RLImage(str(LOGO_PATH), width=2.6 * cm, height=2.6 * cm))
                story.append(Spacer(1, 10))
            except Exception:
                pass
        story.append(Paragraph(_escape_xml(request.title or "Respuesta SHIMIN"), title_style))
        story.append(Paragraph(
            f"SHIMIN Proposal Intelligence · Generado {datetime.now().strftime('%d-%m-%Y %H:%M')}",
            sub_style,
        ))
        story.append(HRFlowable(width="100%", thickness=1.2, color=SHIMIN_COPPER, spaceBefore=4, spaceAfter=12))

        # Cuerpo Markdown completo, sin imprimir sus marcadores literales.
        for block in _parse_markdown(request.answer or ""):
            if block.kind == "heading":
                style = title_style if block.level == 1 else h2_style if block.level == 2 else h3_style
                story.append(Paragraph(_pdf_inline(block.text), style))
            elif block.kind == "bullet":
                story.append(Paragraph(_pdf_inline(block.text), bullet_style, bulletText="•"))
            elif block.kind == "number":
                story.append(Paragraph(_pdf_inline(block.text), bullet_style, bulletText=f"{block.number}."))
            elif block.kind == "quote":
                story.append(Paragraph(_pdf_inline(block.text), quote_style))
            elif block.kind == "code":
                story.append(Paragraph(html_escape(block.text).replace("\n", "<br/>"), code_style))
            elif block.kind == "hr":
                story.append(HRFlowable(width="100%", thickness=0.6, color=SHIMIN_LINE, spaceBefore=5, spaceAfter=7))
            elif block.kind == "table":
                story.append(_pdf_markdown_table(block, body_style))
                story.append(Spacer(1, 8))
            else:
                story.append(Paragraph(_pdf_inline(block.text).replace("\n", "<br/>"), body_style))

        # Tablas
        for table in request.tables[:6]:
            rows = table.get("rows", [])
            if not rows:
                continue
            story.append(Paragraph(_escape_xml(str(table.get("name", "Tabla"))), h2_style))
            columns = list(rows[0].keys())[:6]
            data = [[Paragraph(_pdf_inline(str(c)), table_header_style) for c in columns]]
            for row in rows[:40]:
                data.append([
                    Paragraph(_pdf_inline(_display_value(row.get(col, ""))[:300]), body_style)
                    for col in columns
                ])
            t = Table(data, repeatRows=1, hAlign="LEFT")
            t.setStyle(TableStyle([
                ("BACKGROUND", (0, 0), (-1, 0), SHIMIN_COPPER),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                ("FONTSIZE", (0, 0), (-1, -1), 8.5),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("GRID", (0, 0), (-1, -1), 0.3, SHIMIN_LINE),
                ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, SHIMIN_MIST]),
                ("LEFTPADDING", (0, 0), (-1, -1), 5),
                ("RIGHTPADDING", (0, 0), (-1, -1), 5),
                ("TOPPADDING", (0, 0), (-1, -1), 3),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
            ]))
            story.append(t)
            story.append(Spacer(1, 10))

        # Fuentes
        if request.sources:
            story.append(Paragraph("Fuentes", h2_style))
            for source in request.sources[:20]:
                s = source.model_dump() if hasattr(source, "model_dump") else dict(source)
                kind = _escape_xml(s.get("kind", "src"))
                title = _escape_xml(str(s.get("title", "") or s.get("codigo", ""))[:200])
                codigo = _escape_xml(s.get("codigo", "") or "")
                line = f'<font color="#c8863b"><b>[{kind}]</b></font> '
                if codigo:
                    line += f"<b>{codigo}</b> · "
                line += title
                url = str(s.get("url") or "").strip()
                if _clickable_url(url):
                    line += f' · <link href="{html_escape(url, quote=True)}" color="#1f5f7a"><u>Abrir fuente</u></link>'
                story.append(Paragraph(line, body_style))

        doc.build(story, onFirstPage=_pdf_header_footer, onLaterPages=_pdf_header_footer)
        return path

    def _typst_pdf(self, request: ExportRequest) -> Path:
        typst = self._typst_binary()
        if not typst.exists():
            return self._pdf(request)
        slug = self._slug(request.title)
        typ_path = self.export_dir / f"{slug}.typ"
        pdf_path = self.export_dir / f"{slug}.pdf"
        typ_path.write_text(self._typst_document(request), encoding="utf-8")
        subprocess.run(
            [str(typst), "compile", str(typ_path), str(pdf_path)],
            check=True,
            timeout=90,
            cwd=settings.project_root,
        )
        return pdf_path

    def _typst_document(self, request: ExportRequest) -> str:
        generated_at = datetime.now().strftime("%d-%m-%Y %H:%M")
        parts = [
            '#set document(title: "' + self._typst_inline(request.title) + '")',
            '#let shimin-blue = rgb("0b2239")',
            '#let shimin-cyan = rgb("0097a9")',
            '#let shimin-gray = rgb("f4f6f8")',
            '#set page(',
            '  margin: (left: 1.65cm, right: 1.65cm, top: 2.2cm, bottom: 1.65cm),',
            '  header: rect(width: 100%, height: 1.15cm, fill: shimin-blue, inset: 8pt)[',
            '    #grid(columns: (1fr, auto),',
            '      [#set text(fill: white, size: 10pt); *SHIMIN* | Propuestas y experiencia],',
            '      [#set text(fill: white, size: 8pt); ProyectoHH Agents]',
            '    )',
            '  ],',
            '  footer: grid(columns: (1fr, auto),',
            '    [#set text(size: 7.5pt, fill: rgb("667085")); Documento generado desde Planilla Master / RAG / LLM Wiki],',
            '    [#set text(size: 7.5pt, fill: rgb("667085")); #context counter(page).display("1") ]',
            '  ),',
            ')',
            '#set text(font: "Arial", size: 9.3pt, fill: rgb("202124"))',
            '#show heading: set block(above: 0.8em, below: 0.35em)',
            '#show heading.where(level: 1): set text(size: 18pt, fill: shimin-blue)',
            '#show heading.where(level: 2): set text(size: 13pt, fill: shimin-blue)',
            '#show heading.where(level: 3): set text(size: 10.5pt, fill: rgb("344054"))',
            '#show table.cell: set text(size: 7.0pt)',
            "",
            '#block(fill: shimin-gray, inset: 12pt, radius: 2pt)[',
            "= " + self._typst_inline(request.title),
            "",
            '#text(size: 8pt, fill: rgb("667085"))[' + self._typst_inline(f"Generado: {generated_at}") + "]",
            "]",
            "",
            self._markdownish_to_typst(request.answer),
        ]
        if request.charts:
            parts.append("")
            parts.append("== Graficos")
            for chart in request.charts[:4]:
                parts.append(self._typst_chart(chart))
        for table in request.tables[:8]:
            rows = table.get("rows", [])
            if not rows:
                continue
            parts.append("")
            parts.append("== " + self._typst_inline(str(table.get("name", "Tabla"))))
            parts.append(self._typst_table(rows[:20]))
        if request.sources:
            parts.append("")
            parts.append("== Fuentes")
            for source in request.sources[:20]:
                parts.append(f"- {self._typst_inline(source.kind)}: {self._typst_inline(source.title)}")
        return "\n".join(parts)

    def _typst_table(self, rows: list[dict]) -> str:
        columns = list(rows[0].keys())[:6]
        cells = []
        for column in columns:
            cells.append("[*" + self._typst_inline(str(column)) + "*]")
        for row in rows:
            for column in columns:
                cells.append("[" + self._typst_inline(str(row.get(column, ""))[:260]) + "]")
        return "#table(\n  columns: " + str(len(columns)) + ",\n  inset: 4pt,\n  stroke: 0.35pt + rgb(\"d8dee8\"),\n  " + ",\n  ".join(cells) + ",\n)"

    def _typst_chart(self, chart: dict) -> str:
        rows = chart.get("rows", [])
        if not rows:
            return ""
        x_key = chart.get("x")
        y_key = chart.get("y")
        if chart.get("type") == "line":
            series = chart.get("series") or []
            y_key = series[0] if series else None
        if not x_key or not y_key:
            return ""
        values = [self._number(row.get(y_key)) for row in rows[:12]]
        max_value = max(values) if values else 1
        if max_value <= 0:
            max_value = 1
        bars = []
        for row, value in zip(rows[:12], values):
            label = self._typst_inline(str(row.get(x_key, ""))[:28])
            width = max(3, round((value / max_value) * 100, 2))
            bars.append(
                '#grid(columns: (3.6cm, 1fr, 1.2cm), gutter: 6pt,'
                f'[{label}], [#rect(width: {width}%, height: 7pt, fill: shimin-cyan)], [#align(right)[{value:g}]]'
                ')'
            )
        return "\n".join(
            [
                "=== " + self._typst_inline(str(chart.get("name", "Grafico"))),
                '#block(fill: rgb("fbfcfd"), inset: 8pt, stroke: 0.35pt + rgb("d8dee8"))[',
                *bars,
                "]",
            ]
        )

    def _markdownish_to_typst(self, text: str) -> str:
        lines = []
        for raw in text.splitlines()[:220]:
            line = raw.strip()
            if not line:
                lines.append("")
                continue
            if line.startswith("### "):
                lines.append("=== " + self._typst_inline(line[4:]))
            elif line.startswith("## "):
                lines.append("== " + self._typst_inline(line[3:]))
            elif line.startswith("# "):
                lines.append("= " + self._typst_inline(line[2:]))
            elif line.startswith("- "):
                lines.append("- " + self._typst_inline(line[2:]))
            else:
                lines.append(self._typst_inline(line))
        return "\n".join(lines)

    def _typst_inline(self, value: str) -> str:
        text = str(value or "")
        replacements = {
            "\\": "\\\\",
            "#": "\\#",
            "$": "\\$",
            "%": "\\%",
            "&": "\\&",
            "_": "\\_",
            "*": "\\*",
            "[": "\\[",
            "]": "\\]",
        }
        for old, new in replacements.items():
            text = text.replace(old, new)
        return text

    def _number(self, value: object) -> float:
        try:
            return float(str(value or "0").replace(",", "."))
        except ValueError:
            return 0.0

    def _typst_binary(self) -> Path:
        candidates = [
            settings.project_root / "tools" / "typst" / "typst-x86_64-pc-windows-msvc" / "typst.exe",
            settings.project_root / "tools" / "typst" / "typst",
        ]
        return next((path for path in candidates if path.exists()), candidates[0])

    @property
    def export_dir(self) -> Path:
        return settings.resolve_path(settings.export_dir)

    def _slug(self, value: str) -> str:
        slug = re.sub(r"[^a-zA-Z0-9_-]+", "-", value.strip()).strip("-").lower()
        return (slug or "respuesta")[:70]


# ---- Helpers de formato SHIMIN ----

def _escape_xml(text: str) -> str:
    return (
        str(text or "")
        .replace("&", "&amp;")
        .replace("<", "&lt;")
        .replace(">", "&gt;")
    )


@dataclass
class _MarkdownBlock:
    kind: str
    text: str = ""
    level: int = 0
    number: int = 0
    headers: list[str] = field(default_factory=list)
    rows: list[list[str]] = field(default_factory=list)


_INLINE_RE = re.compile(
    r"(?P<link>\[(?P<link_text>[^\]\n]+)\]\((?P<link_url>https?://[^\s)]+|/api/[^\s)]+)(?:\s+\"[^\"]*\")?\))"
    r"|(?P<code>`(?P<code_text>[^`\n]+)`)"
    r"|(?P<bold>\*\*(?P<bold_text>.+?)\*\*)"
    r"|(?P<bold_u>__(?P<bold_u_text>.+?)__)"
    r"|(?P<strike>~~(?P<strike_text>.+?)~~)"
    r"|(?P<italic>(?<!\*)\*(?P<italic_text>[^*\n]+)\*(?!\*))"
    r"|(?P<italic_u>(?<!\w)_(?P<italic_u_text>[^_\n]+)_(?!\w))"
)


def _parse_markdown(text: str) -> list[_MarkdownBlock]:
    """Convierte Markdown de chat en bloques semánticos compartidos por PDF y Word."""
    lines = str(text or "").replace("\r\n", "\n").replace("\r", "\n").split("\n")
    blocks: list[_MarkdownBlock] = []
    index = 0
    while index < len(lines):
        line = lines[index].rstrip()
        stripped = line.strip()
        if not stripped:
            index += 1
            continue

        fence = re.match(r"^\s*(```+|~~~+)", line)
        if fence:
            marker = fence.group(1)
            index += 1
            code_lines = []
            while index < len(lines) and not re.match(rf"^\s*{re.escape(marker[0])}{{{len(marker)},}}\s*$", lines[index]):
                code_lines.append(lines[index].rstrip())
                index += 1
            if index < len(lines):
                index += 1
            blocks.append(_MarkdownBlock("code", text="\n".join(code_lines)))
            continue

        if _is_markdown_table(lines, index):
            headers = _split_markdown_row(lines[index])
            index += 2
            rows = []
            while index < len(lines) and "|" in lines[index] and lines[index].strip():
                cells = _split_markdown_row(lines[index])
                cells = (cells + [""] * len(headers))[:len(headers)]
                rows.append(cells)
                index += 1
            blocks.append(_MarkdownBlock("table", headers=headers, rows=rows))
            continue

        heading = re.match(r"^\s{0,3}(#{1,6})\s+(.+?)\s*#*\s*$", line)
        if heading:
            blocks.append(_MarkdownBlock("heading", text=heading.group(2), level=len(heading.group(1))))
            index += 1
            continue
        if re.match(r"^\s{0,3}([-*_])(?:\s*\1){2,}\s*$", line):
            blocks.append(_MarkdownBlock("hr"))
            index += 1
            continue
        bullet = re.match(r"^\s*[-+*•]\s+(.+)$", line)
        if bullet:
            blocks.append(_MarkdownBlock("bullet", text=bullet.group(1).strip()))
            index += 1
            continue
        numbered = re.match(r"^\s*(\d+)[.)]\s+(.+)$", line)
        if numbered:
            blocks.append(_MarkdownBlock("number", text=numbered.group(2).strip(), number=int(numbered.group(1))))
            index += 1
            continue
        if re.match(r"^\s*>\s?", line):
            quote_lines = []
            while index < len(lines) and re.match(r"^\s*>\s?", lines[index]):
                quote_lines.append(re.sub(r"^\s*>\s?", "", lines[index]).strip())
                index += 1
            blocks.append(_MarkdownBlock("quote", text=" ".join(quote_lines)))
            continue

        paragraph = [stripped]
        index += 1
        while index < len(lines) and lines[index].strip() and not _starts_markdown_block(lines, index):
            paragraph.append(lines[index].strip())
            index += 1
        blocks.append(_MarkdownBlock("paragraph", text=" ".join(paragraph)))
    return blocks


def _starts_markdown_block(lines: list[str], index: int) -> bool:
    line = lines[index]
    return bool(
        re.match(r"^\s*(```+|~~~+)", line)
        or re.match(r"^\s{0,3}#{1,6}\s+", line)
        or re.match(r"^\s{0,3}([-*_])(?:\s*\1){2,}\s*$", line)
        or re.match(r"^\s*[-+*•]\s+", line)
        or re.match(r"^\s*\d+[.)]\s+", line)
        or re.match(r"^\s*>\s?", line)
        or _is_markdown_table(lines, index)
    )


def _is_markdown_table(lines: list[str], index: int) -> bool:
    if index + 1 >= len(lines) or "|" not in lines[index] or "|" not in lines[index + 1]:
        return False
    headers = _split_markdown_row(lines[index])
    separators = _split_markdown_row(lines[index + 1])
    return bool(headers and len(headers) == len(separators) and all(
        re.fullmatch(r":?-{3,}:?", cell.replace(" ", "")) for cell in separators
    ))


def _split_markdown_row(line: str) -> list[str]:
    value = line.strip().strip("|")
    return [cell.replace(r"\|", "|").strip() for cell in re.split(r"(?<!\\)\|", value)]


def _inline_tokens(text: str) -> list[tuple[str, str, str | None]]:
    tokens: list[tuple[str, str, str | None]] = []
    position = 0
    for match in _INLINE_RE.finditer(str(text or "")):
        if match.start() > position:
            tokens.append(("text", text[position:match.start()], None))
        if match.group("link"):
            tokens.append(("link", match.group("link_text"), match.group("link_url")))
        elif match.group("code"):
            tokens.append(("code", match.group("code_text"), None))
        elif match.group("bold") or match.group("bold_u"):
            tokens.append(("bold", match.group("bold_text") or match.group("bold_u_text"), None))
        elif match.group("strike"):
            tokens.append(("strike", match.group("strike_text"), None))
        else:
            tokens.append(("italic", match.group("italic_text") or match.group("italic_u_text"), None))
        position = match.end()
    if position < len(text):
        tokens.append(("text", text[position:], None))
    return tokens


def _pdf_inline(text: str) -> str:
    parts = []
    for kind, value, url in _inline_tokens(text):
        safe = html_escape(value)
        if kind == "bold":
            parts.append(f"<b>{safe}</b>")
        elif kind == "italic":
            parts.append(f"<i>{safe}</i>")
        elif kind == "strike":
            parts.append(f"<strike>{safe}</strike>")
        elif kind == "code":
            parts.append(f'<font name="Courier">{safe}</font>')
        elif kind == "link" and _clickable_url(url):
            parts.append(f'<link href="{html_escape(str(url), quote=True)}" color="#1f5f7a"><u>{safe}</u></link>')
        else:
            parts.append(safe)
    return "".join(parts)


def _clickable_url(value: str | None) -> bool:
    url = str(value or "").strip()
    return bool(re.match(r"^https?://", url, flags=re.IGNORECASE) or url.startswith("/api/"))


def _display_value(value) -> str:
    if isinstance(value, dict):
        return ", ".join(f"{key}: {_display_value(item)}" for key, item in value.items())
    if isinstance(value, (list, tuple, set)):
        return ", ".join(_display_value(item) for item in value)
    return str(value if value is not None else "")


def _pdf_markdown_table(block: _MarkdownBlock, body_style: ParagraphStyle) -> Table:
    header_style = ParagraphStyle(
        "MarkdownTableHeader", parent=body_style, fontName="Helvetica-Bold",
        fontSize=8.5, leading=10, textColor=colors.white, spaceAfter=0,
    )
    cell_style = ParagraphStyle(
        "MarkdownTableCell", parent=body_style, fontSize=8.5, leading=10, spaceAfter=0,
    )
    data = [[Paragraph(_pdf_inline(header), header_style) for header in block.headers]]
    for row in block.rows[:80]:
        data.append([Paragraph(_pdf_inline(cell), cell_style) for cell in row])
    available_width = 17.4 * cm
    weights = [
        max(6, min(32, max([len(block.headers[index]), *[len(row[index]) for row in block.rows if index < len(row)]], default=6)))
        for index in range(len(block.headers))
    ]
    weight_total = sum(weights) or 1
    widths = [available_width * weight / weight_total for weight in weights]
    table = Table(data, colWidths=widths, repeatRows=1, hAlign="LEFT")
    table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), SHIMIN_COPPER),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("GRID", (0, 0), (-1, -1), 0.3, SHIMIN_LINE),
        ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, SHIMIN_MIST]),
        ("LEFTPADDING", (0, 0), (-1, -1), 5),
        ("RIGHTPADDING", (0, 0), (-1, -1), 5),
        ("TOPPADDING", (0, 0), (-1, -1), 4),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
    ]))
    return table


def _pdf_header_footer(canvas_obj, doc) -> None:
    """Header con logo SHIMIN + footer con número de página y leyenda."""
    canvas_obj.saveState()
    width, height = doc.pagesize
    # Línea cobre superior
    canvas_obj.setStrokeColor(SHIMIN_COPPER)
    canvas_obj.setLineWidth(1.2)
    canvas_obj.line(1.8 * cm, height - 1.3 * cm, width - 1.8 * cm, height - 1.3 * cm)
    # Texto SHIMIN top-right
    canvas_obj.setFont("Helvetica-Bold", 8)
    canvas_obj.setFillColor(SHIMIN_INK)
    canvas_obj.drawRightString(width - 1.8 * cm, height - 1.05 * cm, "SHIMIN · Proposal Intelligence")
    # Footer
    canvas_obj.setFont("Helvetica", 7.5)
    canvas_obj.setFillColor(SHIMIN_MUTED)
    canvas_obj.drawString(1.8 * cm, 1.2 * cm, "Confidencial · Generado desde Planilla Master / RAG / LLM Wiki")
    canvas_obj.drawRightString(width - 1.8 * cm, 1.2 * cm, f"Página {doc.page}")
    canvas_obj.restoreState()


def _docx_horizontal_rule(doc, color: str = "C8863B", size: int = 6) -> None:
    """Añade una línea horizontal en color al documento Word."""
    p = doc.add_paragraph()
    p_pr = p._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def _docx_shade_cell(cell, color: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color)
    tc_pr.append(shd)


def _docx_render_markdown(doc, text: str) -> None:
    """Renderiza los mismos bloques Markdown que el PDF dentro de un DOCX real."""
    for block in _parse_markdown(text):
        if block.kind == "heading":
            paragraph = doc.add_paragraph()
            _docx_apply_inline(paragraph, block.text)
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(16 if block.level == 1 else 14 if block.level == 2 else 11.5)
                run.font.color.rgb = (
                    RGBColor(0x18, 0x2B, 0x36)
                    if block.level == 1 else RGBColor(0x22, 0x3F, 0x4D)
                    if block.level == 2 else RGBColor(0xC8, 0x86, 0x3B)
                )
            paragraph.paragraph_format.space_before = Pt(12 if block.level <= 2 else 8)
            paragraph.paragraph_format.space_after = Pt(4)
        elif block.kind in {"bullet", "number"}:
            paragraph = doc.add_paragraph(style="List Bullet" if block.kind == "bullet" else "List Number")
            _docx_apply_inline(paragraph, block.text)
        elif block.kind == "quote":
            paragraph = doc.add_paragraph()
            paragraph.paragraph_format.left_indent = Cm(0.5)
            paragraph.paragraph_format.right_indent = Cm(0.3)
            _docx_apply_inline(paragraph, block.text)
            for run in paragraph.runs:
                run.italic = True
                run.font.color.rgb = RGBColor(0x5F, 0x74, 0x7D)
            _docx_left_border(paragraph, "C8863B")
        elif block.kind == "code":
            paragraph = doc.add_paragraph()
            paragraph.paragraph_format.left_indent = Cm(0.35)
            run = paragraph.add_run(block.text)
            run.font.name = "Courier New"
            run.font.size = Pt(8.5)
            _docx_shade_paragraph(paragraph, "F4F6F7")
        elif block.kind == "hr":
            _docx_horizontal_rule(doc, color="D7E0E2", size=4)
        elif block.kind == "table":
            _docx_markdown_table(doc, block)
        else:
            paragraph = doc.add_paragraph()
            paragraph.paragraph_format.space_after = Pt(5)
            _docx_apply_inline(paragraph, block.text)


def _docx_apply_inline(paragraph, text: str) -> None:
    for kind, value, url in _inline_tokens(text):
        if kind == "link" and _clickable_url(url):
            _docx_add_hyperlink(paragraph, value, str(url))
            continue
        run = paragraph.add_run(value)
        if kind == "bold":
            run.bold = True
        elif kind == "italic":
            run.italic = True
        elif kind == "strike":
            run.font.strike = True
        elif kind == "code":
            run.font.name = "Courier New"
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0x22, 0x3F, 0x4D)


def _docx_add_hyperlink(paragraph, text: str, url: str) -> None:
    relation_id = paragraph.part.relate_to(url, RT.HYPERLINK, is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relation_id)
    run = OxmlElement("w:r")
    run_properties = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "1F5F7A")
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    run_properties.extend([color, underline])
    run.append(run_properties)
    value = OxmlElement("w:t")
    value.text = text
    run.append(value)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def _docx_markdown_table(doc, block: _MarkdownBlock) -> None:
    table = doc.add_table(rows=1, cols=len(block.headers))
    table.style = "Light Grid Accent 1"
    table.autofit = True
    for index, header in enumerate(block.headers):
        cell = table.rows[0].cells[index]
        cell.text = ""
        _docx_apply_inline(cell.paragraphs[0], header)
        for run in cell.paragraphs[0].runs:
            run.bold = True
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            run.font.size = Pt(9)
        _docx_shade_cell(cell, "C8863B")
    for row_index, values in enumerate(block.rows[:100]):
        cells = table.add_row().cells
        for column_index, value in enumerate(values[:len(block.headers)]):
            cells[column_index].text = ""
            _docx_apply_inline(cells[column_index].paragraphs[0], value)
            for run in cells[column_index].paragraphs[0].runs:
                run.font.size = Pt(8.5)
            if row_index % 2 == 0:
                _docx_shade_cell(cells[column_index], "EEF3F4")


def _docx_shade_paragraph(paragraph, color: str) -> None:
    properties = paragraph._p.get_or_add_pPr()
    shade = OxmlElement("w:shd")
    shade.set(qn("w:fill"), color)
    properties.append(shade)


def _docx_left_border(paragraph, color: str) -> None:
    properties = paragraph._p.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "12")
    left.set(qn("w:space"), "6")
    left.set(qn("w:color"), color)
    borders.append(left)
    properties.append(borders)


# Alias interno para llamadas antiguas dentro de extensiones locales.
def _docx_render_markdown_light(doc, text: str) -> None:
    _docx_render_markdown(doc, text)
