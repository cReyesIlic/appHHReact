"""Drafts de propuestas: el usuario sube antecedentes (PDF/DOCX), el sistema extrae texto,
genera una guía .md con puntos principales y los indexa para que el agente los consuma.

Storage:
  storage/proposal_drafts/<slug>/antecedentes/*.pdf|*.docx
  storage/proposal_drafts/<slug>/texts/*.txt              (texto extraído)
  storage/proposal_drafts/<slug>/guia.md                  (resumen LLM)
  storage/proposal_drafts/<slug>/metadata.json

SQLite:
  proposal_drafts(slug, owner_id, title, cliente, status, created_at, updated_at)
  proposal_draft_files(id, slug, filename, kind, size, chars_extracted, uploaded_at)
  proposal_draft_chunks(id, slug, source_file, text, char_start)
"""

from __future__ import annotations

import json
import re
import shutil
import sqlite3
import time
import unicodedata
import zipfile
from contextlib import closing
from dataclasses import dataclass
from datetime import datetime
from hashlib import sha1
from io import BytesIO
from pathlib import Path

import httpx

from app.core.config import settings


@dataclass
class DraftFileInfo:
    filename: str
    kind: str
    size: int
    chars_extracted: int


class ProposalDraftService:
    def __init__(self) -> None:
        self._last_extraction_method = "none"
        self._ensure_tables()

    # ---- paths ----
    # Los archivos viven en storage/proposal_drafts/<owner_id_safe>/<slug>/...
    # Cada usuario tiene su propio espacio aislado en disco.

    def _owner_root(self, owner_id: str) -> Path:
        safe = re.sub(r"[^A-Za-z0-9._-]+", "_", str(owner_id or "anonymous"))[:80]
        return settings.resolve_path("storage/proposal_drafts") / safe

    def _draft_dir(self, owner_id: str, slug: str) -> Path:
        return self._owner_root(owner_id) / slug

    def _antecedentes_dir(self, owner_id: str, slug: str) -> Path:
        return self._draft_dir(owner_id, slug) / "antecedentes"

    def _texts_dir(self, owner_id: str, slug: str) -> Path:
        return self._draft_dir(owner_id, slug) / "texts"

    def _guia_path(self, owner_id: str, slug: str) -> Path:
        return self._draft_dir(owner_id, slug) / "guia.md"

    def _owner_for_slug(self, slug: str) -> str | None:
        """Resuelve el owner_id de un slug (consultando BD). Usar en métodos que reciben solo slug."""
        with closing(sqlite3.connect(settings.sqlite_path, timeout=5)) as conn, conn:
            row = conn.execute(
                "select owner_id from proposal_drafts where slug = ?", (slug,)
            ).fetchone()
        return row[0] if row else None

    # ---- CRUD ----

    def adopt_aliases(self, canonical_owner_id: str, aliases: tuple[str, ...] | list[str]) -> dict:
        """Mueve borradores de IDs historicos al owner canonico sin perder archivos."""
        legacy = [
            str(value).strip()
            for value in aliases
            if str(value or "").strip() and str(value).strip() != canonical_owner_id
        ]
        if not legacy:
            return {"drafts": 0, "directories": 0, "errors": []}

        placeholders = ",".join("?" for _ in legacy)
        with closing(sqlite3.connect(settings.sqlite_path, timeout=30)) as conn:
            rows = conn.execute(
                f"select slug, owner_id from proposal_drafts where owner_id in ({placeholders})",
                legacy,
            ).fetchall()

        migrated: list[str] = []
        moved = 0
        errors: list[dict] = []
        for slug, old_owner_id in rows:
            source = self._draft_dir(old_owner_id, slug)
            destination = self._draft_dir(canonical_owner_id, slug)
            try:
                if source.exists() and not destination.exists():
                    destination.parent.mkdir(parents=True, exist_ok=True)
                    shutil.move(str(source), str(destination))
                    moved += 1
                elif source.exists() and destination.exists():
                    # Una ejecucion previa pudo quedar a medias. Solo copiamos
                    # archivos ausentes; nunca sobrescribimos una version nueva.
                    for path in source.rglob("*"):
                        relative = path.relative_to(source)
                        target = destination / relative
                        if path.is_dir():
                            target.mkdir(parents=True, exist_ok=True)
                        elif not target.exists():
                            target.parent.mkdir(parents=True, exist_ok=True)
                            shutil.copy2(path, target)
                migrated.append(slug)
            except OSError as exc:
                errors.append({"slug": slug, "error": str(exc)})

        if migrated:
            placeholders = ",".join("?" for _ in migrated)
            with closing(sqlite3.connect(settings.sqlite_path, timeout=30)) as conn, conn:
                conn.execute(
                    f"update proposal_drafts set owner_id = ? where slug in ({placeholders})",
                    (canonical_owner_id, *migrated),
                )
        return {"drafts": len(migrated), "directories": moved, "errors": errors}

    def create_draft(self, owner_id: str, title: str, cliente: str | None = None) -> dict:
        slug = self._slug(title, owner_id)
        now = datetime.now().isoformat(timespec="seconds")
        with closing(sqlite3.connect(settings.sqlite_path, timeout=5)) as conn, conn:
            conn.execute(
                """
                insert or ignore into proposal_drafts
                (slug, owner_id, title, cliente, status, created_at, updated_at)
                values (?, ?, ?, ?, 'draft', ?, ?)
                """,
                (slug, owner_id, title.strip()[:200], (cliente or "").strip()[:120], now, now),
            )
        self._draft_dir(owner_id, slug).mkdir(parents=True, exist_ok=True)
        self._antecedentes_dir(owner_id, slug).mkdir(parents=True, exist_ok=True)
        self._texts_dir(owner_id, slug).mkdir(parents=True, exist_ok=True)
        return self.get_draft(owner_id, slug)

    def list_drafts(self, owner_id: str, limit: int = 100) -> list[dict]:
        with closing(sqlite3.connect(settings.sqlite_path, timeout=5)) as conn, conn:
            rows = conn.execute(
                """
                select slug, title, cliente, status, created_at, updated_at,
                       (select count(*) from proposal_draft_files f where f.slug = d.slug) as files_count
                from proposal_drafts d
                where owner_id = ?
                order by updated_at desc
                limit ?
                """,
                (owner_id, limit),
            ).fetchall()
        return [
            {
                "slug": r[0], "title": r[1], "cliente": r[2], "status": r[3],
                "created_at": r[4], "updated_at": r[5], "files_count": r[6] or 0,
            }
            for r in rows
        ]

    def get_draft(self, owner_id: str, slug: str) -> dict:
        with closing(sqlite3.connect(settings.sqlite_path, timeout=5)) as conn, conn:
            row = conn.execute(
                "select slug, title, cliente, status, created_at, updated_at, "
                "coalesce(brief_text, '') from proposal_drafts "
                "where slug = ? and owner_id = ?",
                (slug, owner_id),
            ).fetchone()
        if not row:
            raise KeyError(slug)
        files = self.list_files(slug)
        guia = self._guia_path(owner_id, slug)
        return {
            "slug": row[0], "title": row[1], "cliente": row[2], "status": row[3],
            "created_at": row[4], "updated_at": row[5],
            "brief_text": row[6],
            "files": files,
            "guide_exists": guia.exists(),
            "guide_path": str(guia) if guia.exists() else None,
        }

    def update_brief(self, owner_id: str, slug: str, brief_text: str) -> dict:
        """Guarda el texto de trabajo del usuario e invalida una guía ya obsoleta."""
        self.get_draft(owner_id, slug)
        brief = str(brief_text or "").strip()[:20000]
        now = datetime.now().isoformat(timespec="seconds")
        with closing(sqlite3.connect(settings.sqlite_path, timeout=5)) as conn, conn:
            conn.execute(
                "update proposal_drafts set brief_text = ?, status = 'draft', updated_at = ? "
                "where slug = ? and owner_id = ?",
                (brief, now, slug, owner_id),
            )
        self._invalidate_guide(owner_id, slug)
        return self.get_draft(owner_id, slug)

    def get_guide(self, owner_id: str, slug: str) -> str:
        path = self._guia_path(owner_id, slug)
        return path.read_text(encoding="utf-8") if path.exists() else ""

    def delete_draft(self, owner_id: str, slug: str) -> dict:
        # Borrar BD primero (rápido)
        with closing(sqlite3.connect(settings.sqlite_path, timeout=5)) as conn, conn:
            n = conn.execute(
                "delete from proposal_drafts where slug = ? and owner_id = ?",
                (slug, owner_id),
            ).rowcount
            conn.execute("delete from proposal_draft_files where slug = ?", (slug,))
            conn.execute("delete from proposal_draft_chunks where slug = ?", (slug,))
        if not n:
            return {"deleted": False}
        # Borrar archivos en disco
        import shutil
        d = self._draft_dir(owner_id, slug)
        if d.exists():
            shutil.rmtree(d, ignore_errors=True)
        return {"deleted": True, "slug": slug}

    # ---- Files ----

    def list_files(self, slug: str) -> list[dict]:
        with closing(sqlite3.connect(settings.sqlite_path, timeout=5)) as conn, conn:
            rows = conn.execute(
                """
                select id, filename, kind, size, chars_extracted, uploaded_at
                from proposal_draft_files where slug = ? order by uploaded_at desc
                """,
                (slug,),
            ).fetchall()
        return [
            {"id": r[0], "filename": r[1], "kind": r[2], "size": r[3],
             "chars_extracted": r[4], "uploaded_at": r[5]}
            for r in rows
        ]

    def add_file(self, owner_id: str, slug: str, filename: str, content: bytes) -> dict:
        """Guarda el archivo, extrae texto, indexa en chunks."""
        # Verificar ownership
        self.get_draft(owner_id, slug)
        clean_name = re.sub(r"[^A-Za-z0-9._-]+", "_", filename)
        kind = self._detect_kind(clean_name)
        if kind not in {"pdf", "docx"}:
            raise ValueError("Solo se aceptan archivos PDF y DOCX")
        self._validate_file_content(content, kind)

        target = self._antecedentes_dir(owner_id, slug) / clean_name
        target.parent.mkdir(parents=True, exist_ok=True)
        target.write_bytes(content)

        # Extraer texto
        text = self._extract_text(content, kind, clean_name)
        text_dir = self._texts_dir(owner_id, slug)
        text_dir.mkdir(parents=True, exist_ok=True)
        text_path = text_dir / f"{Path(clean_name).stem}.txt"
        text_path.write_text(text, encoding="utf-8")

        # Index en chunks
        chunks = self._chunkify(text, max_chars=1600, overlap=200)
        with closing(sqlite3.connect(settings.sqlite_path, timeout=5)) as conn, conn:
            conn.execute(
                """
                insert into proposal_draft_files (slug, filename, kind, size, chars_extracted, uploaded_at)
                values (?, ?, ?, ?, ?, ?)
                on conflict(slug, filename) do update set
                    kind = excluded.kind,
                    size = excluded.size,
                    chars_extracted = excluded.chars_extracted,
                    uploaded_at = excluded.uploaded_at
                """,
                (slug, clean_name, kind, len(content), len(text), datetime.now().isoformat(timespec="seconds")),
            )
            # Reemplazar chunks de este archivo
            conn.execute(
                "delete from proposal_draft_chunks where slug = ? and source_file = ?",
                (slug, clean_name),
            )
            for idx, chunk_text in enumerate(chunks):
                conn.execute(
                    """
                    insert into proposal_draft_chunks (slug, source_file, text, char_start, chunk_index)
                    values (?, ?, ?, ?, ?)
                    """,
                    (slug, clean_name, chunk_text, idx * (1600 - 200), idx),
                )
            conn.execute(
                "update proposal_drafts set status = 'draft', updated_at = ? where slug = ?",
                (datetime.now().isoformat(timespec="seconds"), slug),
            )

        self._invalidate_guide(owner_id, slug)

        return {
            "filename": clean_name,
            "kind": kind,
            "size": len(content),
            "chars_extracted": len(text),
            "chunks_created": len(chunks),
            "extraction_method": self._last_extraction_method,
            "extraction_warning": None if chunks else self._extraction_warning(),
        }

    def delete_file(self, owner_id: str, slug: str, filename: str) -> dict:
        """Elimina un antecedente y todo su texto/chunks, validando ownership."""
        self.get_draft(owner_id, slug)
        clean_name = re.sub(r"[^A-Za-z0-9._-]+", "_", filename)
        with closing(sqlite3.connect(settings.sqlite_path, timeout=30)) as conn:
            exists = conn.execute(
                "select 1 from proposal_draft_files where slug = ? and filename = ?",
                (slug, clean_name),
            ).fetchone()
        if not exists:
            raise FileNotFoundError(clean_name)

        antecedent = self.file_path(owner_id, slug, clean_name)
        extracted = self._texts_dir(owner_id, slug) / f"{Path(clean_name).stem}.txt"
        try:
            antecedent.unlink(missing_ok=True)
            extracted.unlink(missing_ok=True)
        except OSError as exc:
            raise OSError(f"No se pudo eliminar {clean_name} del almacenamiento: {exc}") from exc

        now = datetime.now().isoformat(timespec="seconds")
        with closing(sqlite3.connect(settings.sqlite_path, timeout=30)) as conn, conn:
            files = conn.execute(
                "delete from proposal_draft_files where slug = ? and filename = ?",
                (slug, clean_name),
            ).rowcount
            chunks = conn.execute(
                "delete from proposal_draft_chunks where slug = ? and source_file = ?",
                (slug, clean_name),
            ).rowcount
            conn.execute(
                "update proposal_drafts set status = 'draft', updated_at = ? where slug = ? and owner_id = ?",
                (now, slug, owner_id),
            )
        self._invalidate_guide(owner_id, slug)
        return {"deleted": True, "filename": clean_name, "file_rows": files, "chunks": chunks}

    def reprocess_file(self, owner_id: str, slug: str, filename: str) -> dict:
        """Reextrae e indexa un archivo ya guardado, incluido OCR si corresponde."""
        self.get_draft(owner_id, slug)
        clean_name = re.sub(r"[^A-Za-z0-9._-]+", "_", filename)
        kind = self._detect_kind(clean_name)
        if kind not in {"pdf", "docx"}:
            raise ValueError("Solo se pueden reprocesar archivos PDF y DOCX")
        path = self.file_path(owner_id, slug, clean_name)
        if not path.exists():
            raise FileNotFoundError(clean_name)
        content = path.read_bytes()
        self._validate_file_content(content, kind)
        text = self._extract_text(content, kind, clean_name)
        chunks = self._chunkify(text, max_chars=1600, overlap=200)

        text_dir = self._texts_dir(owner_id, slug)
        text_dir.mkdir(parents=True, exist_ok=True)
        (text_dir / f"{Path(clean_name).stem}.txt").write_text(text, encoding="utf-8")
        now = datetime.now().isoformat(timespec="seconds")
        with closing(sqlite3.connect(settings.sqlite_path, timeout=30)) as conn, conn:
            conn.execute(
                "delete from proposal_draft_chunks where slug = ? and source_file = ?",
                (slug, clean_name),
            )
            for idx, chunk_text in enumerate(chunks):
                conn.execute(
                    """
                    insert into proposal_draft_chunks (slug, source_file, text, char_start, chunk_index)
                    values (?, ?, ?, ?, ?)
                    """,
                    (slug, clean_name, chunk_text, idx * 1400, idx),
                )
            updated = conn.execute(
                """
                update proposal_draft_files
                set kind = ?, size = ?, chars_extracted = ?
                where slug = ? and filename = ?
                """,
                (kind, len(content), len(text), slug, clean_name),
            ).rowcount
            if not updated:
                conn.execute(
                    """
                    insert into proposal_draft_files
                    (slug, filename, kind, size, chars_extracted, uploaded_at)
                    values (?, ?, ?, ?, ?, ?)
                    """,
                    (slug, clean_name, kind, len(content), len(text), now),
                )
            conn.execute(
                "update proposal_drafts set status = 'draft', updated_at = ? where slug = ? and owner_id = ?",
                (now, slug, owner_id),
            )
        self._invalidate_guide(owner_id, slug)
        return {
            "filename": clean_name,
            "kind": kind,
            "size": len(content),
            "chars_extracted": len(text),
            "chunks_created": len(chunks),
            "extraction_method": self._last_extraction_method,
            "extraction_warning": None if chunks else self._extraction_warning(),
        }

    def reprocess_pending(self, owner_id: str, limit: int = 20) -> dict:
        """Reintenta archivos sin texto para corregir fallos transitorios de OCR/indexacion."""
        with closing(sqlite3.connect(settings.sqlite_path, timeout=30)) as conn:
            rows = conn.execute(
                """
                select distinct f.slug, f.filename
                from proposal_draft_files f
                join proposal_drafts d on d.slug = f.slug
                where d.owner_id = ? and coalesce(f.chars_extracted, 0) = 0
                order by f.uploaded_at desc
                limit ?
                """,
                (owner_id, max(1, min(int(limit), 50))),
            ).fetchall()
        repaired: list[dict] = []
        errors: list[dict] = []
        for slug, filename in rows:
            try:
                repaired.append({"slug": slug, **self.reprocess_file(owner_id, slug, filename)})
            except (FileNotFoundError, OSError, ValueError) as exc:
                errors.append({"slug": slug, "filename": filename, "error": str(exc)})
        return {"checked": len(rows), "repaired": repaired, "errors": errors}

    def file_path(self, owner_id: str, slug: str, filename: str) -> Path:
        return self._antecedentes_dir(owner_id, slug) / re.sub(r"[^A-Za-z0-9._-]+", "_", filename)

    async def import_from_sharepoint(
        self,
        owner_id: str,
        slug: str,
        codigo: str,
        filenames: list[str] | None = None,
    ) -> dict:
        """Descarga PDFs/DOCX de la carpeta '01 Informacion Cliente' de una oferta O-XXXX
        y los agrega al draft. Estos son los ANTECEDENTES del cliente (RFP, bases técnicas).
        """
        self.get_draft(owner_id, slug)  # ownership check
        from app.services.sharepoint_client import SharePointClient
        sp = SharePointClient()
        if not sp._configured():
            return {"error": "SharePoint no configurado (faltan TENANT_ID/CLIENT_ID/CLIENT_SECRET)"}
        items = await sp.list_offer_antecedentes(codigo)
        if not items:
            return {
                "codigo": codigo, "found": 0, "imported": 0, "files": [],
                "note": f"No hay archivos en '01 Informacion Cliente' de {codigo} (vacía o no existe).",
            }
        if filenames:
            wanted = [f.lower() for f in filenames]
            items = [i for i in items if any(w in i["name"].lower() for w in wanted)]
        imported: list[dict] = []
        for item in items:
            try:
                content = await sp.download_file(item)
                if not content:
                    continue
                result = self.add_file(owner_id, slug, item["name"], content)
                imported.append({**result, "source": "sharepoint", "codigo_origen": codigo})
            except Exception as exc:  # noqa: BLE001
                imported.append({"filename": item.get("name"), "error": str(exc), "codigo_origen": codigo})
        return {
            "codigo": codigo,
            "found": len(items),
            "imported": sum(1 for f in imported if not f.get("error")),
            "errors": sum(1 for f in imported if f.get("error")),
            "files": imported,
        }

    async def preview_sharepoint(self, codigo: str) -> dict:
        """Lista (sin descargar) los archivos antecedentes disponibles en SharePoint para una oferta."""
        from app.services.sharepoint_client import SharePointClient
        sp = SharePointClient()
        if not sp._configured():
            return {"error": "SharePoint no configurado"}
        items = await sp.list_offer_antecedentes(codigo)
        return {
            "codigo": codigo,
            "count": len(items),
            "files": [
                {"name": i["name"], "size": i.get("size", 0), "kind": i.get("kind"), "webUrl": i.get("webUrl")}
                for i in items
            ],
        }

    # ---- Search dentro del draft ----

    def search_chunks(self, slug: str, query: str, limit: int = 8) -> list[dict]:
        terms = [self._norm(t) for t in (query or "").split() if len(t) >= 3]
        if not terms:
            return []
        with closing(sqlite3.connect(settings.sqlite_path, timeout=5)) as conn, conn:
            rows = conn.execute(
                "select id, source_file, text, char_start from proposal_draft_chunks where slug = ?",
                (slug,),
            ).fetchall()
        hits: list[dict] = []
        for cid, source, text, char_start in rows:
            haystack = self._norm(text)
            score = sum(1 for t in terms if t in haystack)
            if score:
                hits.append({
                    "id": cid, "source_file": source,
                    "char_start": char_start,
                    "snippet": text[:800], "score": float(score),
                })
        hits.sort(key=lambda h: h["score"], reverse=True)
        return hits[:limit]

    def all_text(self, slug: str, max_chars: int = 80000) -> str:
        """Texto completo concatenado (truncado) para enviar al LLM."""
        with closing(sqlite3.connect(settings.sqlite_path, timeout=5)) as conn, conn:
            rows = conn.execute(
                """
                select source_file, text from proposal_draft_chunks
                where slug = ? order by source_file, chunk_index
                """,
                (slug,),
            ).fetchall()
        parts: list[str] = []
        current = ""
        total = 0
        for source, text in rows:
            if source != current:
                parts.append(f"\n\n## Fuente: {source}\n")
                current = source
            parts.append(text)
            total += len(text)
            if total >= max_chars:
                parts.append("\n\n[...truncado...]")
                break
        return "".join(parts)

    def save_guide(self, owner_id: str, slug: str, markdown: str) -> Path:
        path = self._guia_path(owner_id, slug)
        path.parent.mkdir(parents=True, exist_ok=True)
        path.write_text(markdown, encoding="utf-8")
        with closing(sqlite3.connect(settings.sqlite_path, timeout=5)) as conn, conn:
            conn.execute(
                "update proposal_drafts set status = 'guided', updated_at = ? where slug = ? and owner_id = ?",
                (datetime.now().isoformat(timespec="seconds"), slug, owner_id),
            )
        return path

    def _invalidate_guide(self, owner_id: str, slug: str) -> None:
        path = self._guia_path(owner_id, slug)
        if path.exists():
            path.unlink()

    # ---- Internals ----

    def _detect_kind(self, filename: str) -> str:
        lower = filename.lower()
        if lower.endswith(".pdf"):
            return "pdf"
        if lower.endswith(".docx"):
            return "docx"
        return "unknown"

    def _validate_file_content(self, content: bytes, kind: str) -> None:
        if not content:
            raise ValueError("El archivo está vacío; vuelve a descargarlo y súbelo nuevamente")
        if kind == "pdf":
            if content[:1024].find(b"%PDF-") < 0:
                if not content.strip(b"\x00"):
                    raise ValueError(
                        "El PDF recibido contiene solo bytes nulos y está corrupto; "
                        "vuelve a descargar el original y súbelo nuevamente"
                    )
                raise ValueError("El archivo recibido no es un PDF válido (falta la cabecera %PDF-)")
            return
        if kind == "docx":
            try:
                with zipfile.ZipFile(BytesIO(content)) as archive:
                    names = set(archive.namelist())
            except (zipfile.BadZipFile, OSError) as exc:
                raise ValueError("El archivo recibido no es un DOCX válido") from exc
            if "word/document.xml" not in names:
                raise ValueError("El archivo recibido no contiene un documento Word válido")

    def _extract_text(self, content: bytes, kind: str, filename: str) -> str:
        if kind == "pdf":
            local_text = ""
            page_count = 0
            try:
                from PyPDF2 import PdfReader
                reader = PdfReader(BytesIO(content))
                page_count = len(reader.pages)
                pages = []
                for idx, page in enumerate(reader.pages, start=1):
                    try:
                        page_text = (page.extract_text() or "").strip()
                        if page_text:
                            pages.append(f"[página {idx}]\n{page_text}")
                    except Exception:
                        continue
                local_text = "\n\n".join(pages)
                minimum_useful_chars = max(120, page_count * 35)
                if len(local_text) >= minimum_useful_chars:
                    self._last_extraction_method = "pypdf2"
                    return local_text
            except Exception:
                # Un parser local puede fallar en un PDF que Document
                # Intelligence si logra leer. El OCR debe seguir ejecutandose.
                local_text = ""
            ocr_text = self._extract_pdf_with_document_intelligence(content)
            if ocr_text:
                self._last_extraction_method = "document_intelligence"
                return ocr_text
            self._last_extraction_method = "pypdf2_low_text" if local_text else "no_text"
            return local_text
        try:
            if kind == "docx":
                from docx import Document
                doc = Document(BytesIO(content))
                blocks = []
                for para in doc.paragraphs:
                    text = (para.text or "").strip()
                    if text:
                        blocks.append(text)
                for table in doc.tables:
                    for row in table.rows:
                        cells = [c.text.strip() for c in row.cells if c.text.strip()]
                        if cells:
                            blocks.append(" | ".join(cells))
                self._last_extraction_method = "python_docx"
                return "\n".join(blocks)
        except Exception:
            self._last_extraction_method = "error"
            return ""
        self._last_extraction_method = "unsupported"
        return ""

    def _extraction_warning(self) -> str:
        return (
            "El archivo es válido, pero no se pudo extraer texto. "
            "Si está escaneado, revisa la configuración de OCR/Document Intelligence."
        )

    def _extract_pdf_with_document_intelligence(self, content: bytes) -> str:
        endpoint = str(settings.document_intelligence_endpoint or "").strip().rstrip("/")
        key = str(settings.document_intelligence_key or "").strip()
        if not endpoint or not key:
            return ""
        analyze_url = (
            f"{endpoint}/documentintelligence/documentModels/prebuilt-layout:analyze"
            "?api-version=2024-11-30"
        )
        headers = {
            "Ocp-Apim-Subscription-Key": key,
            "Content-Type": "application/pdf",
        }
        try:
            with httpx.Client(timeout=45) as client:
                response = client.post(analyze_url, headers=headers, content=content)
                response.raise_for_status()
                if response.status_code == 200:
                    return self._document_intelligence_text(response.json())
                operation_url = response.headers.get("operation-location")
                if not operation_url:
                    return ""
                poll_headers = {"Ocp-Apim-Subscription-Key": key}
                for _ in range(30):
                    poll = client.get(operation_url, headers=poll_headers)
                    poll.raise_for_status()
                    payload = poll.json()
                    status = str(payload.get("status") or "").casefold()
                    if status == "succeeded":
                        return self._document_intelligence_text(payload)
                    if status in {"failed", "canceled", "cancelled"}:
                        return ""
                    time.sleep(min(float(poll.headers.get("retry-after") or 1), 3.0))
        except (httpx.HTTPError, ValueError, TypeError):
            return ""
        return ""

    def _document_intelligence_text(self, payload: dict) -> str:
        result = payload.get("analyzeResult") or payload.get("analyze_result") or {}
        pages = result.get("pages") or []
        parts: list[str] = []
        for index, page in enumerate(pages, start=1):
            lines = page.get("lines") or []
            text = "\n".join(
                str(line.get("content") or "").strip()
                for line in lines
                if str(line.get("content") or "").strip()
            )
            if text:
                page_number = page.get("pageNumber") or page.get("page_number") or index
                parts.append(f"[página {page_number}]\n{text}")
        if parts:
            return "\n\n".join(parts)
        return str(result.get("content") or "").strip()

    def _chunkify(self, text: str, max_chars: int = 1600, overlap: int = 200) -> list[str]:
        if not text:
            return []
        chunks = []
        start = 0
        while start < len(text):
            piece = text[start:start + max_chars].strip()
            if piece:
                chunks.append(piece)
            start += max_chars - overlap
            if start >= len(text):
                break
        return chunks[:200]  # safety cap

    def _slug(self, title: str, owner_id: str) -> str:
        base = re.sub(r"[^a-z0-9-]+", "-", self._norm(title)).strip("-")[:50] or "draft"
        suffix = sha1(f"{owner_id}:{title}:{datetime.now().isoformat()}".encode()).hexdigest()[:6]
        return f"{base}-{suffix}"

    def _norm(self, value: object) -> str:
        text = unicodedata.normalize("NFKD", str(value or "").lower())
        return "".join(ch for ch in text if not unicodedata.combining(ch))

    def _ensure_tables(self) -> None:
        settings.sqlite_path.parent.mkdir(parents=True, exist_ok=True)
        with closing(sqlite3.connect(settings.sqlite_path)) as conn, conn:
            conn.execute(
                """
                create table if not exists proposal_drafts (
                    slug text primary key,
                    owner_id text not null,
                    title text not null,
                    cliente text,
                    brief_text text default '',
                    status text default 'draft',
                    created_at text not null,
                    updated_at text not null
                )
                """
            )
            draft_columns = {
                row[1] for row in conn.execute("pragma table_info(proposal_drafts)").fetchall()
            }
            if "brief_text" not in draft_columns:
                conn.execute("alter table proposal_drafts add column brief_text text default ''")
            conn.execute(
                """
                create table if not exists proposal_draft_files (
                    id integer primary key autoincrement,
                    slug text not null,
                    filename text not null,
                    kind text not null,
                    size integer not null,
                    chars_extracted integer default 0,
                    uploaded_at text not null
                )
                """
            )
            # Versiones anteriores insertaban una fila nueva al volver a subir
            # el mismo nombre. Conservamos la más reciente y prevenimos dobles.
            conn.execute(
                """
                delete from proposal_draft_files
                where id not in (
                    select max(id) from proposal_draft_files group by slug, filename
                )
                """
            )
            conn.execute(
                """
                create table if not exists proposal_draft_chunks (
                    id integer primary key autoincrement,
                    slug text not null,
                    source_file text not null,
                    text text not null,
                    char_start integer,
                    chunk_index integer
                )
                """
            )
            conn.execute("create index if not exists idx_drafts_owner on proposal_drafts(owner_id, updated_at)")
            conn.execute("create index if not exists idx_draft_files_slug on proposal_draft_files(slug)")
            conn.execute(
                "create unique index if not exists idx_draft_files_unique "
                "on proposal_draft_files(slug, filename)"
            )
            conn.execute("create index if not exists idx_draft_chunks_slug on proposal_draft_chunks(slug, source_file)")
